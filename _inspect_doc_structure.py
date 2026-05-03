import sys
from docx import Document

path = sys.argv[1]
d = Document(path)
print('FILE', path)
print('PARA_COUNT', len(d.paragraphs))
print('TABLE_COUNT', len(d.tables))
for i, para in enumerate(d.paragraphs, 1):
    t = (para.text or '').strip()
    if t:
        print(f"{i:03d}|{para.style.name}|{t[:220]}")
