from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path(r"c:\Users\VEDANT\OneDrive\Desktop\mini project")
TARGETS = [
    ROOT / "miniproject.docx",
    ROOT / "Mini Project.with_diagrams.docx",
]


def p_text(p) -> str:
    return (p.text or "").strip()


def find_idx(doc: Document, text: str, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if p_text(doc.paragraphs[i]) == text:
            return i
    return -1


def remove_between(doc: Document, start_heading: str, end_heading: str):
    s = find_idx(doc, start_heading)
    if s < 0:
        return
    e = find_idx(doc, end_heading, s + 1)
    if e < 0 or e <= s:
        return

    # Remove paragraphs strictly between the two headings.
    to_remove = doc.paragraphs[s + 1 : e]
    for p in to_remove:
        el = p._element
        el.getparent().remove(el)


def insert_lines_before_heading(doc: Document, end_heading: str, lines: Iterable[str]):
    e = find_idx(doc, end_heading)
    if e < 0:
        return
    anchor = doc.paragraphs[e]

    for line in lines:
        if line == "__PAGE_BREAK__":
            p = anchor.insert_paragraph_before("")
            r = p.add_run()
            r.add_break(WD_BREAK.PAGE)
            continue
        anchor.insert_paragraph_before(line)


def replace_section(doc: Document, start_heading: str, end_heading: str, lines: list[str]):
    remove_between(doc, start_heading, end_heading)
    insert_lines_before_heading(doc, end_heading, lines)


def replace_paragraph_exact(doc: Document, old: str, new: str):
    for p in doc.paragraphs:
        if p_text(p) == old:
            p.text = new
            return


def refresh_doc(path: Path) -> Path:
    doc = Document(str(path))

    # Top matter normalization.
    replace_paragraph_exact(doc, "CuraVision AI Ã¢â‚¬â€œ AI Powered Medical Image Classification System", "CuraVision AI - AI Powered Medical Image Classification System")
    replace_paragraph_exact(doc, "MINI PROJECT Ã¢â‚¬â€œ BCA20090", "MINI PROJECT - BCA20090")
    replace_paragraph_exact(doc, "Date: 1st February 2026", "Date: 23 April 2026")

    # Index refresh.
    replace_section(
        doc,
        "Index",
        "Nomenclature/Notations",
        [
            "1. Nomenclature/Notations",
            "2. Abstract",
            "3. Motivation",
            "4. Literature Review",
            "5. Problem Formulation and Objectives",
            "6. Methodology / Planning of Work",
            "   6.1 Requirement Analysis",
            "   6.2 Implementation Phases",
            "   6.3 System Architecture",
            "   6.4 Database Design",
            "   6.5 Database Tables",
            "   6.6 Project Flow Diagram",
            "   6.7 Testing and Deployment",
            "7. Dataset Specifications",
            "8. Testing and Evaluation",
            "   8.1 Input and Output Screenshots (Blank Page)",
            "9. Facilities Required for Proposed Work",
            "10. References",
            "",
        ],
    )

    # Abstract.
    replace_section(
        doc,
        "Abstract",
        "Motivation",
        [
            "CuraVision AI is an educational clinical decision support system for chest X-ray triage that combines secure doctor access, two-stage AI inference, case persistence, and patient report lookup.",
            "The deployed workflow now reflects real project updates: doctor registration/login with OTP verification, X-ray upload with optional patient photo, CLIP-based chest suitability validation, TensorFlow pneumonia classification, and AI-generated insight summaries.",
            "Predictions, confidence values, uploaded image metadata, and report history are stored in a SQLite database through Flask-SQLAlchemy, enabling dashboard analytics and longitudinal case review.",
            "The system is designed for academic demonstration only and supports clinician-in-the-loop usage with transparent outputs and clear medical disclaimers.",
            "",
        ],
    )

    # Motivation.
    replace_section(
        doc,
        "Motivation",
        "Literature Review",
        [
            "Pneumonia diagnosis support is highly valuable in settings where specialist radiology access is limited and quick triage is required.",
            "A major weakness in many student AI prototypes is direct classification of any uploaded image without suitability checks, producing unreliable results for non-medical inputs.",
            "CuraVision AI addresses this gap by enforcing a two-stage pipeline: image suitability validation first, medical classification second.",
            "The project motivation also includes full-stack learning outcomes: authentication, secure report handling, persistence, AI inference, and clinically readable output design in one integrated platform.",
            "",
        ],
    )

    # Literature review summary text.
    replace_section(
        doc,
        "Literature Review",
        "Problem Formulation & Objectives",
        [
            "Existing medical image analysis approaches range from manual radiologist interpretation to single-stage deep learning classifiers.",
            "Manual review offers clinical reliability but can be slow and resource intensive; direct one-stage AI can be fast but may fail when non-chest or out-of-distribution images are submitted.",
            "The reviewed trend supports hybrid systems that combine validation, inference, and workflow integration, which aligns with the updated CuraVision AI architecture.",
            "",
        ],
    )

    # Problem and objectives refresh.
    replace_section(
        doc,
        "Problem Formulation & Objectives",
        "Methodology/ Planning of Work",
        [
            "Problem Statement and Objectives",
            "The project addresses reliability and workflow gaps in educational medical AI systems by integrating secure user flow, suitability validation, and persistent case management.",
            "Problem Statement",
            "Conventional classroom models may produce predictions for irrelevant images and often lack proper doctor workflow, case history, and database-backed auditability.",
            "Main Objective",
            "Build an AI-assisted web system that validates chest X-rays, predicts NORMAL/PNEUMONIA with confidence, and stores complete case records for review.",
            "Specific Objectives",
            "Implement two-stage inference with CLIP validation followed by TensorFlow classification.",
            "Integrate doctor authentication with OTP verification and session-controlled access.",
            "Persist case records (patient details, files, predictions, confidence, AI insights) in SQLite.",
            "Provide dashboard, case history, patient report portal, and helper recommendations.",
            "",
        ],
    )

    # Section 6 updates.
    replace_section(
        doc,
        "6.1 Requirement Analysis",
        "6.2 Implementation Phases",
        [
            "The updated requirement analysis reflects both model and platform evolution.",
            "Functional Requirements",
            "Secure Access: Doctors must register and verify email OTP before accessing clinical workflows.",
            "Upload and Validation: System must accept chest X-ray uploads and reject non-chest inputs via suitability detection.",
            "Prediction and Explanation: System must return NORMAL/PNEUMONIA prediction, confidence, and AI-generated insight text.",
            "Case Persistence: Each prediction event must be stored with patient metadata and timestamps for audit and follow-up.",
            "Patient Report Access: Patient identifier-based report viewing should be available for retrieval of prior reports.",
            "Software Requirements",
            "Python 3.11, Flask, Flask-SQLAlchemy, Flask-Bcrypt, TensorFlow/Keras, Transformers (CLIP), Pillow/OpenCV.",
            "Hardware Requirements",
            "Minimum 8 GB RAM, multicore CPU, and optional CUDA GPU for faster model experimentation.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.2 Implementation Phases",
        "6.3 System Architecture",
        [
            "System Design: Modular architecture with auth, routes, AI logic, database models, and templates.",
            "Frontend Development: Dashboard, upload form, case list/detail, AI helper, and patient portal templates.",
            "Backend Development: Flask blueprints for auth and main workflow, robust validation and file handling.",
            "Database Integration: SQLAlchemy models for doctors and patient case records with schema upgrade checks.",
            "Model Integration: Startup loading of pneumonia classifier and CLIP suitability model for real-time inference.",
            "Testing and Hardening: Route-level checks, validation tests, and usability walkthroughs across core workflows.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.3 System Architecture",
        "6.4 Database Design",
        [
            "The current architecture follows a secure web inference pattern.",
            "Access Layer: Doctor registration, OTP verification, login, and protected dashboard routes.",
            "Input Layer: Upload endpoint accepts X-ray image with optional patient photo and structured patient metadata.",
            "Validation Layer: CLIP suitability check filters non-chest uploads before medical classification.",
            "Inference Layer: TensorFlow model predicts NORMAL/PNEUMONIA using configured threshold logic.",
            "Insight Layer: Rule-based helper generates recommendation and summary text.",
            "Persistence Layer: SQLite stores doctors, cases, blobs/paths, predictions, confidence, and timestamps.",
            "Presentation Layer: Dashboard and report pages expose operational and patient-facing outputs.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.5 Database Tables",
        "6.6 Project Flow Diagram",
        [
            "Doctors Table (doctors)",
            "Columns: id, name, email, password_hash, is_verified, verification_otp, otp_expires_at.",
            "Purpose: Doctor authentication, OTP verification lifecycle, and controlled access.",
            "Patients Table (patients)",
            "Columns: id, patient_identifier, name, age, symptoms, generated_symptoms, image_path, photo_path, image_blob, image_mime, photo_blob, photo_mime, prediction, confidence, ai_insights, created_at.",
            "Purpose: Persistent case records for prediction history, detail view, and patient report lookup.",
            "Database Operations",
            "Create/Read: New records on prediction, ordered retrieval for dashboard/case listings.",
            "Delete: Case deletion removes both DB record and associated uploaded files safely.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.7 Testing and Deployment",
        "Dataset Specifications",
        [
            "Development Phases and Deployment",
            "Environment Setup: Virtual environment, dependency installation, and model artifact setup.",
            "Run Strategy: Application starts directly using saved model files, avoiding retraining during server startup.",
            "API Integration: Upload, prediction, report retrieval, and helper routes validated end-to-end.",
            "Database Migration Safety: Startup checks add missing columns for older SQLite files without data loss.",
            "Deployment Scope: Local deployment for academic demonstration with clear responsible-AI disclaimers.",
            "",
        ],
    )

    # Dataset section refresh.
    replace_section(
        doc,
        "7. Dataset Specifications",
        "Testing and Evaluation",
        [
            "CuraVision AI uses both ML training datasets and runtime application data.",
            "Primary Classification Dataset",
            "Source: Kaggle Chest X-ray Pneumonia dataset (NORMAL and PNEUMONIA classes).",
            "Usage: Training/evaluation of the TensorFlow pneumonia classifier.",
            "Detector Dataset",
            "Source Structure: detector_dataset with chest_xray and not_chest_xray classes.",
            "Usage: Supporting suitability gating logic and rejection behavior for non-chest images.",
            "Runtime Application Dataset (Database-backed)",
            "Storage: SQLite database plus uploads directory for xray/photo assets.",
            "Data Captured: patient id, demographics, symptoms, image metadata/blobs, prediction, confidence, insights, timestamp.",
            "7.1 Data Attributes and Metadata",
            "Case_ID / Patient Identifier: Tracks case history and patient report access path.",
            "Prediction Fields: NORMAL/PNEUMONIA output, confidence score, and helper-generated insights.",
            "Media Fields: File paths plus optional blob copies and MIME types for robust report rendering.",
            "7.2 Dataset Integrity and Validation",
            "Input Validation: Extension and format checks on upload.",
            "Clinical Relevance Validation: Chest suitability screening before classification.",
            "Preprocessing Consistency: Resize and normalization pipeline aligned to model requirements.",
            "",
        ],
    )

    # Testing section refresh + blank screenshot page.
    replace_section(
        doc,
        "8. Testing and Evaluations",
        "Facilities Required for Proposed Work",
        [
            "To validate the updated full-stack system, the following testing tracks are used:",
            "Functional Testing: Registration/login/OTP flow, upload process, prediction JSON/UI rendering, case list/detail, and patient reports.",
            "Validation Testing: Suitability detector behavior for chest vs non-chest uploads and failure messaging.",
            "Model Testing: Prediction confidence behavior, threshold handling, and class output consistency.",
            "Database Testing: Record creation, retrieval ordering, and case deletion impacts on file/database state.",
            "Usability Testing: Dashboard readability, report navigation, and helper interpretation by end users.",
            "8.1 Input and Output Screenshots",
            "__PAGE_BREAK__",
            "INPUT AND OUTPUT SCREENSHOTS (BLANK PAGE)",
            "Insert screenshots for the following here:",
            "1) Input Upload Screen",
            "2) Prediction Result Screen",
            "3) Case Detail Screen",
            "4) Patient Report Screen",
            "",
            "",
            "",
        ],
    )

    # Facilities refresh.
    replace_section(
        doc,
        "Facilities Required for Proposed Work",
        "Bibliography/References",
        [
            "The following facilities are required for development, execution, and demonstration of the updated system:",
            "9.1 Hardware Requirements:",
            "Development Machine: Minimum 8 GB RAM (16 GB recommended), modern multicore CPU.",
            "Optional GPU: CUDA-capable GPU for faster experimentation/training iterations.",
            "Storage and Backup: Space for dataset, model files, and SQLite/application backup snapshots.",
            "9.2 Software Requirements:",
            "Operating System: Windows/macOS/Linux.",
            "IDE and Tooling: VS Code, Git, GitHub.",
            "Backend Stack: Python, Flask, Flask-SQLAlchemy, Flask-Bcrypt.",
            "AI Stack: TensorFlow/Keras, Transformers CLIP, Pillow/OpenCV.",
            "Database: SQLite with schema compatibility checks on startup.",
            "",
        ],
    )

    # Reference heading normalization and content refresh.
    replace_paragraph_exact(doc, "8. Testing and Evaluations", "8. Testing and Evaluation")
    replace_paragraph_exact(doc, "10. References", "10. References")

    # Replace detailed references list.
    start_ref = find_idx(doc, "10. References")
    if start_ref >= 0:
        # Remove everything after 10. References and rebuild reference list.
        for p in doc.paragraphs[start_ref + 1 :]:
            el = p._element
            el.getparent().remove(el)

        anchor_end = doc.paragraphs[-1]
        ref_lines = [
            "1. Kaggle Chest X-ray Pneumonia Dataset.",
            "2. TensorFlow and Keras official documentation.",
            "3. Hugging Face Transformers and CLIP documentation.",
            "4. Flask and Flask-SQLAlchemy official documentation.",
            "5. MobileNetV2 research paper and transfer learning references.",
            "6. SQLite documentation for lightweight relational storage.",
            "",
        ]
        for line in ref_lines:
            anchor_end.insert_paragraph_before(line)

    try:
        doc.save(str(path))
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}.updated_content{path.suffix}")
        doc.save(str(fallback))
        return fallback


if __name__ == "__main__":
    for target in TARGETS:
        if not target.exists():
            print(f"Skipped missing file: {target}")
            continue
        out = refresh_doc(target)
        print(f"Updated: {out}")
