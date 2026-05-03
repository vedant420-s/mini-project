from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"c:\Users\VEDANT\OneDrive\Desktop\mini project")
DESKTOP = Path(r"c:\Users\VEDANT\OneDrive\Desktop")
DIAGRAM_DIR = ROOT / "doc_diagrams" / "professional_faculty_v2"
OUTPUT_DOC = DESKTOP / "AI_Medical_Image_Classifier_Faculty_Report.docx"

PROJECT_TITLE = "CuraVision AI - AI Powered Medical Image Classification System"
TEAM_MEMBERS = [
    "Vedant Khedkar - 1272240644",
    "Sakshi Pandey - 1272240631",
]
GUIDES = [
    "Dr. Pradeep Kumar Tiwari",
    "Mrs. Aparna Kulkarni",
]
REPO = "https://github.com/vedant420-s/mini-project"

COLOR_NAVY = RGBColor(23, 54, 93)
COLOR_TEAL = RGBColor(21, 125, 132)
COLOR_GRAY = RGBColor(88, 96, 108)


def set_page_border(section):
    sect_pr = section._sectPr
    borders = OxmlElement("w:pgBorders")
    borders.set(qn("w:offsetFrom"), "page")

    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "24")
        elem.set(qn("w:color"), "264478")
        borders.append(elem)

    sect_pr.append(borders)


def style_heading(paragraph, color=COLOR_NAVY):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = "Cambria"
        run.font.bold = True
        run.font.color.rgb = color


def add_para(doc: Document, text: str, size: int = 12, line_spacing: float = 1.3):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = line_spacing
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(size)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(12)


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Number")
    p.paragraph_format.space_after = Pt(5)
    for r in p.runs:
        r.font.name = "Calibri"
        r.font.size = Pt(12)


def add_caption(doc: Document, caption: str):
    p = doc.add_paragraph(caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = COLOR_GRAY


def new_major_section(doc: Document, title: str):
    doc.add_page_break()
    h = doc.add_heading(title, level=1)
    style_heading(h)


def diagram_path(name: str) -> Path:
    p = DIAGRAM_DIR / name
    if not p.exists():
        raise FileNotFoundError(f"Required diagram not found: {p}")
    return p


def add_placeholder_box(doc: Document, label: str):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    row = table.rows[0]
    row.height = Inches(3.0)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    cell = table.cell(0, 0)
    cell.text = "\n\nSCREENSHOT PLACEHOLDER\n" + label + "\n(Insert final working screenshot here)"
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(12)
            r.font.bold = True
            r.font.color.rgb = COLOR_GRAY


def build_report():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    set_page_border(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(12)

    title = doc.add_paragraph("MINI PROJECT - BCA20090")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.runs[0]
    r.font.name = "Cambria"
    r.font.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COLOR_NAVY

    for line in [
        "SY BCA 2025-26",
        "SYNOPSIS REPORT",
        "TITLE:",
        PROJECT_TITLE,
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Cambria" if line in ["SYNOPSIS REPORT", PROJECT_TITLE] else "Calibri"
            run.font.bold = line in ["SYNOPSIS REPORT", PROJECT_TITLE]
            run.font.size = Pt(15 if line == PROJECT_TITLE else 13)
            run.font.color.rgb = COLOR_NAVY

    doc.add_paragraph("")
    p = doc.add_paragraph("Submitted by:")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.bold = True
    for member in TEAM_MEMBERS:
        add_bullet(doc, member)

    add_para(doc, "Guided By:")
    for g in GUIDES:
        add_para(doc, g)

    add_para(doc, "Dr. Vishwanath Karad MIT World Peace University")
    add_para(doc, "Department of Computer Science and Applications")
    add_para(doc, "Academic Year: 2025-26")
    add_para(doc, f"Date: {datetime.now().strftime('%d %B %Y')}")
    add_para(doc, f"Repository: {REPO}")

    doc.add_page_break()

    h = doc.add_heading("Index", level=1)
    style_heading(h)
    index_items = [
        "1. Nomenclature/Notations",
        "2. Abstract",
        "3. Motivation",
        "4. Literature Review",
        "5. Problem Formulation & Objectives",
        "6. Methodology/ Planning of Work",
        "6.1 Requirement Analysis",
        "6.2 Implementation Phases",
        "6.3 System Architecture",
        "6.4 Database Design",
        "6.5 Database Tables",
        "6.6 Project Flow Diagram",
        "6.7 Testing and Deployment",
        "7. Dataset Specifications",
        "8. Testing and Evaluation",
        "8.1 Input and Output Screenshots",
        "9. Facilities Required for Proposed Work",
        "10. References",
    ]
    for item in index_items:
        add_number(doc, item)

    new_major_section(doc, "Nomenclature/Notations")
    nomenclature = [
        "Artificial Intelligence (AI): Systems capable of performing learning and decision-support tasks from data.",
        "Clinical Decision Support System (CDSS): Software that assists clinicians by presenting evidence-based insights.",
        "Convolutional Neural Network (CNN): Deep-learning architecture specialized for image feature extraction.",
        "Chest X-ray (CXR): Radiographic image modality used to assess thoracic conditions such as pneumonia.",
        "Out-of-Distribution (OOD): Inputs that differ from model training domain and may produce unreliable outputs.",
        "Transfer Learning: Adapting a pre-trained model to a related target problem.",
        "MobileNetV2: Efficient CNN backbone used in many low-resource computer vision tasks.",
        "Confidence Score: Probability-linked value indicating prediction certainty.",
        "ORM: Object Relational Mapping layer used to map Python classes to database tables.",
        "OTP: One-time password used for email-based verification of user ownership.",
    ]
    for term in nomenclature:
        add_bullet(doc, term)

    new_major_section(doc, "Abstract")
    add_para(doc, "CuraVision AI is an educational medical image analysis project designed to assist chest X-ray triage in a structured and responsible workflow. The system integrates authentication, image validation, inference, result presentation, and persistent storage in a single web-based platform.")
    add_para(doc, "The core pipeline uses two stages. Stage one checks whether uploaded content is a valid chest X-ray to reduce misuse and improve workflow reliability. Stage two classifies validated images into NORMAL or PNEUMONIA and returns confidence values, explanatory text, and timestamped output for review.")
    add_para(doc, "Beyond model prediction, the application supports clinical-style case management by storing patient identifiers, symptoms, generated summaries, and report history. This enables dashboard-level tracking and case-wise retrieval, which is important for academic demonstration of real-world software behavior.")
    add_para(doc, "The project is not intended for direct clinical deployment. It is presented as a learning-oriented decision-support prototype where every output must be validated by qualified healthcare professionals.")

    new_major_section(doc, "Motivation")
    add_para(doc, "Pneumonia remains a major respiratory concern, and timely triage is often critical. In settings with limited specialist availability, preliminary AI-assisted screening can support faster prioritization and improve workflow continuity.")
    add_para(doc, "Traditional student-level model demonstrations typically focus only on classification accuracy and skip practical safeguards. Such systems may generate medical outputs for invalid images and fail to provide traceable case records. This gap motivated the design of CuraVision AI as a full-stack, validation-first workflow.")
    add_para(doc, "The project also serves an academic purpose: to demonstrate integration across machine learning, backend engineering, database modeling, and human-centered interface design. By combining these layers, the report reflects not just an algorithm but a complete engineered system.")
    add_para(doc, "An additional motivation is ethical responsibility. The solution explicitly includes confidence visibility, report persistence, and medical disclaimers to reduce over-trust and encourage clinician-in-the-loop interpretation.")

    new_major_section(doc, "Literature Review")
    add_para(doc, "Published chest X-ray classification studies show that transfer learning can deliver strong baseline performance even with moderate computational resources. Lightweight backbones such as MobileNetV2 are frequently used in educational and edge scenarios.")
    add_para(doc, "Recent applied AI workflows also emphasize pre-validation and data-quality gating before medical inference. These practices reduce out-of-domain risk and improve robustness in practical deployments.")
    add_para(doc, "Based on this review, CuraVision AI combines established classification techniques with workflow controls including suitability checks, structured case records, and explainable confidence outputs.")

    review_table = doc.add_table(rows=5, cols=4)
    review_table.style = "Table Grid"
    review_table.cell(0, 0).text = "Approach"
    review_table.cell(0, 1).text = "Strength"
    review_table.cell(0, 2).text = "Limitation"
    review_table.cell(0, 3).text = "Relevance to Project"
    rows = [
        ("Manual Radiology Review", "High expert reliability", "Time-intensive and specialist dependent", "Used as final validation benchmark"),
        ("Single-stage CNN", "Fast inference", "May mis-handle OOD inputs", "Improved with validation-first design"),
        ("Transfer Learning MobileNetV2", "Efficient and accurate baseline", "Needs domain-aware preprocessing", "Adopted as classification backbone"),
        ("Two-stage Validation + Classification", "Higher practical robustness", "More system complexity", "Core architecture of CuraVision AI"),
    ]
    for i, row in enumerate(rows, start=1):
        for j, cell in enumerate(row):
            review_table.cell(i, j).text = cell

    new_major_section(doc, "Problem Formulation & Objectives")
    add_para(doc, "Problem Statement: Existing academic AI demos often classify any uploaded image as a medical class without checking relevance, which creates unreliable outputs and weak trust in decision-support contexts.")
    add_para(doc, "Main Objective: Develop a reliable AI-powered system that validates chest X-ray inputs, predicts pneumonia risk, and stores complete case records for review and reporting.")
    add_para(doc, "Specific Objectives:")
    objectives = [
        "Implement two-stage inference with suitability gating followed by pneumonia classification.",
        "Expose confidence values and explanatory summaries for interpretation transparency.",
        "Integrate doctor authentication and OTP verification for controlled system access.",
        "Store patient metadata, model outputs, and media references in persistent database tables.",
        "Provide patient identifier-based report retrieval workflow.",
        "Maintain clear educational and ethical boundaries for usage context.",
    ]
    for obj in objectives:
        add_bullet(doc, obj)

    new_major_section(doc, "Methodology/ Planning of Work")
    add_para(doc, "The project follows a staged engineering methodology: requirement definition, modular implementation, controlled testing, and documentation for reproducible academic evaluation.")
    add_para(doc, "A layered architecture was chosen to isolate concerns: authentication, workflow routes, AI inference logic, persistence, and presentation. This approach improves maintainability and allows focused validation of each subsystem.")
    add_para(doc, "Planning also considered operational readiness, including environment setup, model artifact handling, and schema compatibility during runtime initialization.")

    new_major_section(doc, "6.1 Requirement Analysis")
    add_para(doc, "Functional requirements were mapped directly to implemented features to ensure traceability from problem statement to working modules.")
    add_para(doc, "Core functional requirements include secure doctor onboarding, X-ray upload, suitability filtering, classification, report persistence, dashboard analytics, and patient-side report lookup.")
    add_para(doc, "Non-functional requirements focus on clarity, reliability, and safe failure behavior. Validation messages are explicit, route access is controlled by sessions, and database operations include rollback paths during exceptions.")
    add_para(doc, "Technology requirements include Python 3.11, Flask ecosystem packages, TensorFlow/Keras, Transformers CLIP, SQLite persistence, and frontend templates rendered through Jinja.")

    new_major_section(doc, "6.2 Implementation Phases")
    phase_lines = [
        "Phase 1 - Environment and Dependency Setup: Python environment creation, package installation, and model artifact placement.",
        "Phase 2 - Data and Model Preparation: Dataset structure validation, preprocessing consistency checks, and loading of trained models.",
        "Phase 3 - Backend Workflow Construction: Authentication routes, prediction API, case management endpoints, and validation routines.",
        "Phase 4 - Frontend Integration: Dashboard, upload, case detail, and patient report templates connected to backend logic.",
        "Phase 5 - Persistence and Media Handling: SQLAlchemy models, blob/path storage strategy, and safe deletion behavior.",
        "Phase 6 - Testing and Refinement: Route checks, manual user-journey validation, and error-path stabilization.",
    ]
    for line in phase_lines:
        add_bullet(doc, line)

    add_para(doc, "Each phase was validated before moving forward, reducing integration defects and ensuring that the final application remained stable under common usage scenarios.")

    new_major_section(doc, "6.3 System Architecture")
    add_para(doc, "The system uses a client-server architecture with modular backend services. User-facing pages interact with Flask routes, which invoke inference components and database models to complete case workflows.")
    add_para(doc, "Authentication and authorization are enforced before accessing clinical routes. Prediction processing remains isolated from presentation logic to keep the pipeline testable and maintainable.")
    doc.add_picture(str(diagram_path("system_architecture.png")), width=Inches(6.8))
    add_caption(doc, "Figure 6.3.1: System architecture")

    new_major_section(doc, "6.4 Database Design")
    add_para(doc, "Database design is centered on two entities: doctors and patients. Doctor records support secure access and verification lifecycle. Patient records represent scan-level reports with metadata, predictions, confidence, and generated insights.")
    add_para(doc, "The design supports both file-path and blob storage options for media retrieval. This dual approach improves robustness for case review screens and minimizes dependency on a single storage path.")
    add_para(doc, "Schema compatibility is maintained through startup checks that add missing fields when required, supporting incremental evolution of the project database during development.")
    doc.add_picture(str(diagram_path("er_diagram.png")), width=Inches(6.8))
    add_caption(doc, "Figure 6.4.1: ER diagram and database relationship")

    new_major_section(doc, "6.5 Database Tables")
    add_para(doc, "Doctors Table (doctors):")
    doctors_table = doc.add_table(rows=8, cols=2)
    doctors_table.style = "Table Grid"
    doctors_table.cell(0, 0).text = "Column"
    doctors_table.cell(0, 1).text = "Description"
    doctors_rows = [
        ("id", "Primary key"),
        ("name", "Doctor full name"),
        ("email", "Unique indexed login email"),
        ("password_hash", "Hashed password"),
        ("is_verified", "OTP verification state"),
        ("verification_otp", "Temporary OTP value"),
        ("otp_expires_at", "OTP expiry timestamp"),
    ]
    for i, (c, d) in enumerate(doctors_rows, start=1):
        doctors_table.cell(i, 0).text = c
        doctors_table.cell(i, 1).text = d

    add_para(doc, "Patients Table (patients):")
    patient_table = doc.add_table(rows=13, cols=2)
    patient_table.style = "Table Grid"
    patient_table.cell(0, 0).text = "Column"
    patient_table.cell(0, 1).text = "Description"
    patient_rows = [
        ("id", "Primary key"),
        ("patient_identifier", "Patient lookup ID"),
        ("name, age, symptoms", "Patient and clinical input fields"),
        ("generated_symptoms", "Auto-generated symptom hints"),
        ("image_path/photo_path", "Stored file references"),
        ("image_blob/photo_blob", "Binary media backup fields"),
        ("image_mime/photo_mime", "Media content types"),
        ("prediction", "Model class output"),
        ("confidence", "Prediction confidence percentage"),
        ("ai_insights", "Rule-based explanation summary"),
        ("created_at", "Case creation timestamp"),
        ("(derived route linkage)", "Connected to case and patient report views"),
    ]
    for i, (c, d) in enumerate(patient_rows, start=1):
        patient_table.cell(i, 0).text = c
        patient_table.cell(i, 1).text = d

    new_major_section(doc, "6.6 Project Flow Diagram")
    add_para(doc, "The workflow covers doctor onboarding, secure session access, case upload, two-stage inference, report persistence, and patient lookup. This end-to-end sequence reflects how the implemented components interact during real use.")
    doc.add_picture(str(diagram_path("clinical_workflow.png")), width=Inches(6.8))
    add_caption(doc, "Figure 6.6.1: End-to-end project flow")

    new_major_section(doc, "6.7 Testing and Deployment")
    add_para(doc, "Testing was performed at three levels: feature validation, inference validation, and data persistence consistency. Route behavior, error handling, and report accessibility were verified through both script checks and manual walkthroughs.")
    add_para(doc, "Deployment for evaluation is local, using Python runtime and pre-trained model artifacts. The runtime setup prioritizes reproducibility for faculty demonstration over cloud-scale optimization.")
    add_para(doc, "The inference pipeline diagram below summarizes processing from input to output for test and deployment understanding.")
    doc.add_picture(str(diagram_path("ai_pipeline.png")), width=Inches(6.8))
    add_caption(doc, "Figure 6.7.1: AI inference pipeline")

    new_major_section(doc, "7. Dataset Specifications")
    add_para(doc, "Primary model training relies on the Kaggle chest X-ray pneumonia dataset. Runtime behavior additionally uses application-generated case records stored in SQLite.")
    add_para(doc, "Dataset-level attributes include class labels, image dimensions, and class imbalance profile. Preprocessing enforces resize and normalization to align with trained model expectations.")
    add_para(doc, "Application-level dataset attributes include patient identifiers, symptom text, prediction outputs, confidence values, generated insights, and timestamps. This allows report traceability across sessions.")
    add_para(doc, "Integrity constraints include strict patient identifier normalization, file-format checks, and suitability filtering prior to classification. Together these controls improve robustness of recorded outcomes.")

    new_major_section(doc, "8. Testing and Evaluation")
    add_para(doc, "Functional testing covered registration, OTP verification, login, upload, prediction rendering, case listing, case detail retrieval, patient report lookup, and case deletion workflows.")
    add_para(doc, "Validation testing focused on incorrect input handling, including missing fields, invalid age values, unsupported file types, and non-chest image submissions. Error responses were checked for clarity and safety.")
    add_para(doc, "Model evaluation references documented metrics from project artifacts. Reported values include training accuracy 92.64%, validation accuracy 87.66%, and test accuracy 87.66%, with high pneumonia recall indicating strong triage sensitivity in this educational setup.")

    eval_table = doc.add_table(rows=6, cols=2)
    eval_table.style = "Table Grid"
    eval_table.cell(0, 0).text = "Metric"
    eval_table.cell(0, 1).text = "Value"
    metrics = [
        ("Training Accuracy", "92.64%"),
        ("Validation Accuracy", "87.66%"),
        ("Test Accuracy", "87.66%"),
        ("NORMAL F1-Score", "81.53%"),
        ("PNEUMONIA F1-Score", "90.73%"),
    ]
    for i, (k, v) in enumerate(metrics, start=1):
        eval_table.cell(i, 0).text = k
        eval_table.cell(i, 1).text = v

    new_major_section(doc, "8.1 Input and Output Screenshots")
    add_para(doc, "Reserved spaces below should be replaced with final working screenshots from your project execution.")
    add_placeholder_box(doc, "Input Upload Screen")
    add_caption(doc, "Figure 8.1.1: Input upload screen")
    add_placeholder_box(doc, "Prediction Result Screen")
    add_caption(doc, "Figure 8.1.2: Prediction result screen")

    doc.add_page_break()
    h = doc.add_heading("8.1 Input and Output Screenshots (Continued)", level=2)
    style_heading(h, COLOR_TEAL)
    add_placeholder_box(doc, "Case Detail Screen")
    add_caption(doc, "Figure 8.1.3: Case detail screen")
    add_placeholder_box(doc, "Patient Report Screen")
    add_caption(doc, "Figure 8.1.4: Patient report screen")

    new_major_section(doc, "9. Facilities Required for Proposed Work")
    add_para(doc, "Hardware Requirements:")
    hardware = [
        "Development machine with minimum 8 GB RAM (16 GB recommended).",
        "Modern multicore CPU for stable local inference and web serving.",
        "Optional CUDA-capable GPU for faster training experimentation.",
        "Adequate storage for datasets, model artifacts, and report backups.",
    ]
    for line in hardware:
        add_bullet(doc, line)

    add_para(doc, "Software Requirements:")
    software = [
        "Operating system: Windows, macOS, or Linux.",
        "IDE and tools: Visual Studio Code, Git, GitHub.",
        "Backend stack: Python, Flask, Flask-SQLAlchemy, Flask-Bcrypt.",
        "AI stack: TensorFlow/Keras, Transformers CLIP, Pillow/OpenCV.",
        "Database: SQLite with startup compatibility checks.",
    ]
    for line in software:
        add_bullet(doc, line)

    new_major_section(doc, "10. References")
    references = [
        "Kaggle Chest X-ray Pneumonia Dataset: https://www.kaggle.com/paultimothymooney/chest-xray-pneumonia",
        "TensorFlow Documentation: https://www.tensorflow.org/",
        "Keras Documentation: https://keras.io/",
        "Flask Documentation: https://flask.palletsprojects.com/",
        "Flask-SQLAlchemy Documentation: https://flask-sqlalchemy.palletsprojects.com/",
        "Transformers Documentation: https://huggingface.co/docs/transformers/",
        "MobileNetV2 Paper: https://arxiv.org/abs/1801.04381",
        "SQLite Documentation: https://www.sqlite.org/docs.html",
        "Project Repository: https://github.com/vedant420-s/mini-project",
    ]
    for ref in references:
        add_bullet(doc, ref)

    doc.save(str(OUTPUT_DOC))


def main():
    build_report()
    print(f"Generated guideline-aligned report: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
