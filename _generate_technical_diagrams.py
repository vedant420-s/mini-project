"""
Generate 4 Technical Diagrams for CuraVision AI:
1. ER Diagram (Entity Relationship)
2. Data Flow Diagram (DFD)
3. Flowchart (System Processes)
4. Database Working Diagram (Schema & Operations)
"""

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Color Palette
NAVY = (23, 61, 93)
TEAL = (21, 125, 132)
ORANGE = (211, 123, 53)
GREEN = (62, 144, 91)
GRAY = (90, 98, 112)
LIGHT_GRAY = (220, 225, 230)
WHITE = (255, 255, 255)
BLACK = (0, 0, 0)

DIAGRAM_WIDTH = 2400
DIAGRAM_HEIGHT = 1500

def draw_box(draw, x, y, width, height, text, fill_color, text_color=WHITE, font_size=14):
    """Draw a rounded rectangle with text"""
    # Draw box
    draw.rectangle([x, y, x + width, y + height], fill=fill_color, outline=BLACK, width=2)
    
    # Draw text
    try:
        font = ImageFont.truetype("arial.ttf", font_size)
    except:
        font = ImageFont.load_default()
    
    text_bbox = draw.textbbox((0, 0), text, font=font)
    text_width = text_bbox[2] - text_bbox[0]
    text_height = text_bbox[3] - text_bbox[1]
    
    text_x = x + (width - text_width) // 2
    text_y = y + (height - text_height) // 2
    
    draw.text((text_x, text_y), text, fill=text_color, font=font)

def draw_arrow(draw, x1, y1, x2, y2, color=BLACK, width=3):
    """Draw an arrow from (x1, y1) to (x2, y2)"""
    draw.line([(x1, y1), (x2, y2)], fill=color, width=width)
    
    # Arrow head
    angle = 20
    import math
    dx = x2 - x1
    dy = y2 - y1
    length = math.sqrt(dx*dx + dy*dy)
    if length > 0:
        dx /= length
        dy /= length
        
        p1x = x2 - 20 * dx - 10 * (-dy)
        p1y = y2 - 20 * dy - 10 * dx
        p2x = x2 - 20 * dx + 10 * (-dy)
        p2y = y2 - 20 * dy + 10 * dx
        
        draw.polygon([(x2, y2), (p1x, p1y), (p2x, p2y)], fill=color)

def draw_label(draw, x, y, text, color=BLACK, font_size=12, bold=False):
    """Draw text label"""
    try:
        font = ImageFont.truetype("arial.ttf", font_size)
    except:
        font = ImageFont.load_default()
    draw.text((x, y), text, fill=color, font=font)

# ==================== ER DIAGRAM ====================
def create_er_diagram():
    """Create Entity Relationship Diagram"""
    img = Image.new('RGB', (DIAGRAM_WIDTH, DIAGRAM_HEIGHT), WHITE)
    draw = ImageDraw.Draw(img)
    
    # Title
    try:
        title_font = ImageFont.truetype("arial.ttf", 28)
    except:
        title_font = ImageFont.load_default()
    draw.text((DIAGRAM_WIDTH//2 - 200, 50), "Entity Relationship Diagram", fill=NAVY, font=title_font)
    
    # Doctor Entity
    doctor_x, doctor_y = 200, 250
    draw.rectangle([doctor_x, doctor_y, doctor_x + 350, doctor_y + 400], 
                   fill=TEAL, outline=BLACK, width=3)
    draw_label(draw, doctor_x + 100, doctor_y + 20, "DOCTOR", color=WHITE, font_size=18)
    
    doctor_fields = [
        "PK: id (INT)",
        "name (VARCHAR)",
        "email (VARCHAR)",
        "password_hash (VARCHAR)",
        "is_verified (BOOL)",
        "verification_otp (VARCHAR)",
        "otp_expires_at (DATETIME)"
    ]
    
    y_offset = doctor_y + 60
    for field in doctor_fields:
        draw_label(draw, doctor_x + 20, y_offset, field, color=WHITE, font_size=11)
        y_offset += 45
    
    # Patient Entity
    patient_x, patient_y = 1850, 250
    draw.rectangle([patient_x, patient_y, patient_x + 350, patient_y + 600], 
                   fill=ORANGE, outline=BLACK, width=3)
    draw_label(draw, patient_x + 80, patient_y + 20, "PATIENT", color=WHITE, font_size=18)
    
    patient_fields = [
        "PK: id (INT)",
        "patient_identifier (VARCHAR)",
        "name (VARCHAR)",
        "age (INT)",
        "symptoms (TEXT)",
        "generated_symptoms (TEXT)",
        "image_path (VARCHAR)",
        "photo_path (VARCHAR)",
        "image_blob (BLOB)",
        "image_mime (VARCHAR)",
        "photo_blob (BLOB)",
        "photo_mime (VARCHAR)",
        "prediction (VARCHAR)",
        "confidence (FLOAT)",
        "ai_insights (TEXT)",
        "created_at (DATETIME)",
        "FK: doctor_id (INT)"
    ]
    
    y_offset = patient_y + 60
    for field in patient_fields:
        draw_label(draw, patient_x + 15, y_offset, field, color=WHITE, font_size=10)
        y_offset += 32
    
    # Relationship
    rel_x = 600
    rel_y = 400
    draw.rectangle([rel_x, rel_y, rel_x + 220, rel_y + 100], 
                   fill=GREEN, outline=BLACK, width=3)
    draw_label(draw, rel_x + 30, rel_y + 25, "examines", color=WHITE, font_size=14)
    draw_label(draw, rel_x + 35, rel_y + 55, "1 : N", color=WHITE, font_size=12)
    
    # Relationship arrows
    draw_arrow(draw, doctor_x + 350, 450, rel_x, 450, NAVY, 4)
    draw_arrow(draw, rel_x + 220, 450, patient_x, 450, NAVY, 4)
    
    # Cardinality labels
    draw_label(draw, 580, 420, "1", color=NAVY, font_size=14)
    draw_label(draw, 1820, 420, "N", color=NAVY, font_size=14)
    
    img.save("doc_diagrams/er_diagram.png")
    return "doc_diagrams/er_diagram.png"

# ==================== DATA FLOW DIAGRAM ====================
def create_dfd():
    """Create Data Flow Diagram"""
    img = Image.new('RGB', (DIAGRAM_WIDTH, DIAGRAM_HEIGHT), WHITE)
    draw = ImageDraw.Draw(img)
    
    # Title
    try:
        title_font = ImageFont.truetype("arial.ttf", 28)
    except:
        title_font = ImageFont.load_default()
    draw.text((DIAGRAM_WIDTH//2 - 250, 50), "Data Flow Diagram (DFD) - Level 1", fill=NAVY, font=title_font)
    
    # Processes (circles)
    def draw_circle(x, y, radius, text, color):
        draw.ellipse([x-radius, y-radius, x+radius, y+radius], fill=color, outline=BLACK, width=3)
        try:
            font = ImageFont.truetype("arial.ttf", 12)
        except:
            font = ImageFont.load_default()
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        draw.text((x - text_width//2, y - text_height//2), text, fill=WHITE, font=font)
    
    # External entities (rectangles)
    draw_box(draw, 100, 300, 200, 120, "Doctor/Patient", NAVY)
    draw_box(draw, DIAGRAM_WIDTH - 300, 300, 200, 120, "Database\n(SQLite)", TEAL)
    
    # Processes
    y_centers = [500, 850, 1100]
    process_texts = ["Upload & Validation", "AI Prediction", "Store & Report"]
    
    for i, (y, text) in enumerate(zip(y_centers, process_texts)):
        draw_circle(800, y, 80, f"{i+1}", ORANGE if i % 2 == 0 else GREEN)
        draw_label(draw, 920, y - 20, text, color=BLACK, font_size=13)
    
    # Data flows
    # Input
    draw_arrow(draw, 300, 360, 720, 500, TEAL, 4)
    draw_label(draw, 350, 320, "X-ray Image", color=TEAL, font_size=11)
    
    # P1 to P2
    draw_arrow(draw, 800, 580, 800, 770, ORANGE, 4)
    draw_label(draw, 820, 675, "Valid Image", color=ORANGE, font_size=11)
    
    # P2 to P3
    draw_arrow(draw, 800, 930, 800, 1020, GREEN, 4)
    draw_label(draw, 820, 975, "Prediction", color=GREEN, font_size=11)
    
    # P3 to DB
    draw_arrow(draw, 880, 1100, DIAGRAM_WIDTH - 300, 360, NAVY, 4)
    draw_label(draw, 1400, 700, "Save Record", color=NAVY, font_size=11)
    
    # DB to Output
    draw_arrow(draw, DIAGRAM_WIDTH - 300, 410, 1050, 850, TEAL, 4)
    draw_label(draw, 1600, 600, "Report Data", color=TEAL, font_size=11)
    
    # Store legend
    draw_label(draw, 100, 1350, "P1: Upload & Validation (CLIP image check)", color=BLACK, font_size=11)
    draw_label(draw, 100, 1380, "P2: AI Prediction (MobileNetV2 CNN inference)", color=BLACK, font_size=11)
    draw_label(draw, 100, 1410, "P3: Store & Report (Database storage + Report generation)", color=BLACK, font_size=11)
    
    img.save("doc_diagrams/dfd_diagram.png")
    return "doc_diagrams/dfd_diagram.png"

# ==================== FLOWCHART ====================
def create_flowchart():
    """Create System Flowchart"""
    img = Image.new('RGB', (DIAGRAM_WIDTH, DIAGRAM_HEIGHT), WHITE)
    draw = ImageDraw.Draw(img)
    
    # Title
    try:
        title_font = ImageFont.truetype("arial.ttf", 28)
    except:
        title_font = ImageFont.load_default()
    draw.text((DIAGRAM_WIDTH//2 - 250, 50), "System Flowchart - Doctor Workflow", fill=NAVY, font=title_font)
    
    # Start
    draw_circle_flowchart = lambda x, y, text, color: draw.ellipse(
        [x-50, y-40, x+50, y+40], fill=color, outline=BLACK, width=2)
    
    def draw_oval(x, y, w, h, text, color):
        draw.ellipse([x-w, y-h, x+w, y+h], fill=color, outline=BLACK, width=2)
        try:
            font = ImageFont.truetype("arial.ttf", 11)
        except:
            font = ImageFont.load_default()
        text_bbox = draw.textbbox((0, 0), text, font=font)
        text_width = text_bbox[2] - text_bbox[0]
        text_height = text_bbox[3] - text_bbox[1]
        draw.text((x - text_width//2, y - text_height//2), text, fill=WHITE, font=font)
    
    def draw_diamond(x, y, size, text, color):
        points = [(x, y - size), (x + size, y), (x, y + size), (x - size, y)]
        draw.polygon(points, fill=color, outline=BLACK, width=2)
        try:
            font = ImageFont.truetype("arial.ttf", 11)
        except:
            font = ImageFont.load_default()
        text_bbox = draw.textbbox((0, 0), text, font=font)
        draw.text((x - (text_bbox[2] - text_bbox[0])//2, y - (text_bbox[3] - text_bbox[1])//2), 
                  text, fill=WHITE, font=font)
    
    # Flow elements
    x_center = 200
    y_pos = 150
    
    # Start
    draw_oval(x_center, y_pos, 60, 35, "START", GREEN)
    y_pos += 100
    draw_arrow(draw, x_center, y_pos - 50, x_center, y_pos - 20, BLACK, 3)
    
    # Register
    draw_box(draw, x_center - 80, y_pos - 20, 160, 80, "Doctor\nRegistration", TEAL)
    y_pos += 110
    draw_arrow(draw, x_center, y_pos - 20, x_center, y_pos, BLACK, 3)
    
    # Email Verification
    draw_box(draw, x_center - 80, y_pos, 160, 80, "Email\nVerification (OTP)", ORANGE)
    y_pos += 110
    draw_arrow(draw, x_center, y_pos - 20, x_center, y_pos, BLACK, 3)
    
    # Login
    draw_box(draw, x_center - 80, y_pos, 160, 80, "Login", TEAL)
    y_pos += 110
    draw_arrow(draw, x_center, y_pos - 20, x_center, y_pos, BLACK, 3)
    
    # Dashboard
    draw_box(draw, x_center - 80, y_pos, 160, 80, "View Dashboard", NAVY, LIGHT_GRAY)
    y_pos += 110
    draw_arrow(draw, x_center, y_pos - 20, x_center, y_pos, BLACK, 3)
    
    # Decision: Upload X-ray?
    draw_diamond(x_center, y_pos + 60, 70, "Upload\nX-ray?", GRAY)
    
    # Yes path
    draw_arrow(draw, x_center + 70, y_pos + 60, x_center + 350, y_pos + 60, BLACK, 3)
    draw_label(draw, x_center + 150, y_pos + 30, "YES", color=BLACK, font_size=11)
    
    # Upload process
    draw_box(draw, x_center + 270, y_pos + 20, 160, 80, "Upload X-ray\nImage", TEAL)
    draw_arrow(draw, x_center + 430, y_pos + 60, x_center + 430, y_pos + 130, BLACK, 3)
    
    # Validation
    draw_box(draw, x_center + 270, y_pos + 130, 160, 80, "CLIP Validation\n(Check if X-ray)", ORANGE)
    draw_arrow(draw, x_center + 430, y_pos + 190, x_center + 430, y_pos + 260, BLACK, 3)
    
    # Prediction
    draw_box(draw, x_center + 270, y_pos + 260, 160, 80, "MobileNetV2\nPrediction", GREEN)
    draw_arrow(draw, x_center + 430, y_pos + 320, x_center + 430, y_pos + 390, BLACK, 3)
    
    # Store
    draw_box(draw, x_center + 270, y_pos + 390, 160, 80, "Store Result\nin Database", TEAL)
    
    # No path (from decision)
    draw_arrow(draw, x_center, y_pos + 130, x_center, y_pos + 260, BLACK, 3)
    draw_label(draw, x_center - 40, y_pos + 180, "NO", color=BLACK, font_size=11)
    
    # View Cases
    draw_box(draw, x_center - 80, y_pos + 260, 160, 80, "View Previous\nCases", NAVY, LIGHT_GRAY)
    
    # End
    draw_arrow(draw, x_center, y_pos + 340, x_center, y_pos + 420, BLACK, 3)
    draw_oval(x_center, y_pos + 470, 60, 35, "END", RED := (220, 50, 50))
    
    img.save("doc_diagrams/flowchart.png")
    return "doc_diagrams/flowchart.png"

# ==================== DATABASE WORKING DIAGRAM ====================
def create_db_working_diagram():
    """Create Database Working Diagram"""
    img = Image.new('RGB', (DIAGRAM_WIDTH, DIAGRAM_HEIGHT), WHITE)
    draw = ImageDraw.Draw(img)
    
    # Title
    try:
        title_font = ImageFont.truetype("arial.ttf", 28)
    except:
        title_font = ImageFont.load_default()
    draw.text((DIAGRAM_WIDTH//2 - 350, 50), "Database Working Diagram - Schema & Operations", 
              fill=NAVY, font=title_font)
    
    # Doctor Table
    draw_box(draw, 100, 200, 450, 500, "", TEAL)
    draw_label(draw, 150, 220, "DOCTOR TABLE", color=WHITE, font_size=16)
    draw.line([(100, 260), (550, 260)], fill=WHITE, width=2)
    
    doc_fields = [
        "id (PK)",
        "name",
        "email (UNIQUE)",
        "password_hash",
        "is_verified",
        "verification_otp",
        "otp_expires_at"
    ]
    
    y_offset = 280
    for field in doc_fields:
        draw_label(draw, 120, y_offset, field, color=WHITE, font_size=12)
        y_offset += 30
    
    # Patient Table
    draw_box(draw, 1350, 200, 450, 650, "", ORANGE)
    draw_label(draw, 1400, 220, "PATIENT TABLE", color=WHITE, font_size=16)
    draw.line([(1350, 260), (1800, 260)], fill=WHITE, width=2)
    
    patient_fields = [
        "id (PK)",
        "patient_identifier",
        "name",
        "age",
        "symptoms",
        "generated_symptoms",
        "image_path",
        "photo_path",
        "image_blob",
        "image_mime",
        "photo_blob",
        "photo_mime",
        "prediction",
        "confidence",
        "ai_insights",
        "created_at",
        "doctor_id (FK)"
    ]
    
    y_offset = 280
    for field in patient_fields:
        draw_label(draw, 1370, y_offset, field, color=WHITE, font_size=11)
        y_offset += 24
    
    # Relationship arrow
    draw_arrow(draw, 550, 450, 1350, 550, NAVY, 5)
    draw_label(draw, 800, 400, "1 Doctor : N Patients", color=NAVY, font_size=12)
    
    # Operations section
    ops_y = 1100
    draw_label(draw, 100, ops_y, "KEY DATABASE OPERATIONS:", color=NAVY, font_size=14)
    
    operations = [
        "1. CREATE: Doctor registers → New row in DOCTOR table with OTP",
        "2. UPDATE: Doctor verified → is_verified flag updated",
        "3. INSERT: X-ray uploaded → New PATIENT record inserted",
        "4. READ: Retrieve patient case history → Query PATIENT where doctor_id = X",
        "5. UPDATE: Prediction complete → Update prediction, confidence, ai_insights",
        "6. DELETE: (Controlled) Remove patient record if needed"
    ]
    
    for op in operations:
        draw_label(draw, 100, ops_y + 40, op, color=BLACK, font_size=11)
        ops_y += 35
    
    img.save("doc_diagrams/db_working_diagram.png")
    return "doc_diagrams/db_working_diagram.png"

# ==================== MAIN DOCUMENT CREATION ====================
def create_technical_diagrams_document():
    """Create Word document with all 4 diagrams"""
    
    # Create diagrams folder if not exists
    os.makedirs("doc_diagrams", exist_ok=True)
    
    print("Generating ER Diagram...")
    er_path = create_er_diagram()
    
    print("Generating Data Flow Diagram...")
    dfd_path = create_dfd()
    
    print("Generating Flowchart...")
    flowchart_path = create_flowchart()
    
    print("Generating Database Working Diagram...")
    db_path = create_db_working_diagram()
    
    # Create Word document
    doc = Document()
    
    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CuraVision AI")
    title_run.font.size = Pt(48)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(*NAVY)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Technical Diagrams & System Architecture")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.color.rgb = RGBColor(*TEAL)
    
    doc.add_paragraph()
    
    project_info = doc.add_paragraph()
    project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = project_info.add_run(
        "Chest X-ray Pneumonia Classification System\n"
        "AI Medical Image Classifier\n"
        "Team: Vedant Khedkar & Sakshi Pandey"
    )
    info_run.font.size = Pt(12)
    
    # Page break
    doc.add_page_break()
    
    # ER Diagram
    h1 = doc.add_heading("1. Entity Relationship (ER) Diagram", 1)
    h1.runs[0].font.color.rgb = RGBColor(*NAVY)
    
    p = doc.add_paragraph(
        "The ER diagram represents the database schema showing the relationship between DOCTOR and PATIENT entities. "
        "Each doctor can examine multiple patients (1:N relationship). The diagram displays all attributes including "
        "primary keys (PK) and foreign keys (FK) that maintain referential integrity."
    )
    p.runs[0].font.size = Pt(11)
    
    doc.add_picture(er_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph("Figure 1: Entity Relationship Diagram")
    caption.runs[0].font.italic = True
    caption.runs[0].font.size = Pt(10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # DFD
    h2 = doc.add_heading("2. Data Flow Diagram (DFD)", 1)
    h2.runs[0].font.color.rgb = RGBColor(*NAVY)
    
    p = doc.add_paragraph(
        "The DFD illustrates how data flows through the system across three main processes: "
        "(1) Upload & Validation - receives X-ray images and validates using CLIP model, "
        "(2) AI Prediction - performs inference using MobileNetV2 CNN, "
        "(3) Store & Report - persists results to database and generates reports. "
        "External entities (Doctor/Patient and Database) interact with these processes."
    )
    p.runs[0].font.size = Pt(11)
    
    doc.add_picture(dfd_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph("Figure 2: Data Flow Diagram (Level 1)")
    caption.runs[0].font.italic = True
    caption.runs[0].font.size = Pt(10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Flowchart
    h3 = doc.add_heading("3. System Flowchart", 1)
    h3.runs[0].font.color.rgb = RGBColor(*NAVY)
    
    p = doc.add_paragraph(
        "The flowchart depicts the complete doctor workflow within the system. Starting from registration and email verification, "
        "doctors can log in, access the dashboard, and choose to either upload new X-ray images for analysis or review previous cases. "
        "The upload path includes CLIP validation to verify the image is a valid X-ray, followed by MobileNetV2 prediction and database storage. "
        "Decision points allow for branching logic based on user actions."
    )
    p.runs[0].font.size = Pt(11)
    
    doc.add_picture(flowchart_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph("Figure 3: Doctor Workflow Flowchart")
    caption.runs[0].font.italic = True
    caption.runs[0].font.size = Pt(10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Database Working Diagram
    h4 = doc.add_heading("4. Database Working Diagram", 1)
    h4.runs[0].font.color.rgb = RGBColor(*NAVY)
    
    p = doc.add_paragraph(
        "This diagram shows the database schema in detail with all tables, attributes, and relationships. "
        "The DOCTOR table stores authentication information and verification status. "
        "The PATIENT table maintains comprehensive medical records including image data, AI predictions, confidence scores, and AI-generated insights. "
        "Key operations including CREATE, READ, UPDATE, and DELETE demonstrate how data is manipulated throughout the application lifecycle."
    )
    p.runs[0].font.size = Pt(11)
    
    doc.add_picture(db_path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    caption = doc.add_paragraph("Figure 4: Database Schema & Operations")
    caption.runs[0].font.italic = True
    caption.runs[0].font.size = Pt(10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    output_path = r"C:\Users\VEDANT\OneDrive\Desktop\CuraVision_Technical_Diagrams.docx"
    doc.save(output_path)
    
    print(f"\n✓ Technical diagrams document created successfully!")
    print(f"✓ Saved to: {output_path}")
    print(f"\nDiagrams included:")
    print(f"  1. ER Diagram - Database schema and relationships")
    print(f"  2. Data Flow Diagram - System data movement")
    print(f"  3. Flowchart - Doctor workflow and processes")
    print(f"  4. Database Working Diagram - Schema with operations")

if __name__ == "__main__":
    create_technical_diagrams_document()
