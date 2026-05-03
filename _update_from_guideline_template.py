from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


GUIDELINE_DOC = Path(r"C:\Users\VEDANT\Downloads\Guidelines_for_PBL_report (1).docx")
OUT_DESKTOP = Path(r"C:\Users\VEDANT\OneDrive\Desktop\PBL.docx")
OUT_WORKSPACE = Path(r"C:\Users\VEDANT\OneDrive\Desktop\mini project\PBL.docx")

DIAGRAM_DIR = Path(r"C:\Users\VEDANT\OneDrive\Desktop\mini project\doc_diagrams\professional_faculty_v2")
DIAGRAMS = {
    "er": DIAGRAM_DIR / "er_diagram.png",
    "flow": DIAGRAM_DIR / "clinical_workflow.png",
    "architecture": DIAGRAM_DIR / "system_architecture.png",
    "pipeline": DIAGRAM_DIR / "ai_pipeline.png",
}


def p_text(p) -> str:
    return (p.text or "").strip()


def find_idx(doc: Document, text: str, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if p_text(doc.paragraphs[i]) == text:
            return i
    return -1


def remove_paragraph(p):
    el = p._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_between(doc: Document, start_text: str, end_text: str) -> bool:
    s = find_idx(doc, start_text)
    if s < 0:
        return False
    e = find_idx(doc, end_text, s + 1)
    if e < 0 or e <= s:
        return False

    for p in doc.paragraphs[s + 1 : e]:
        remove_paragraph(p)
    return True


def insert_lines_before(doc: Document, anchor_text: str, lines: Iterable[str]) -> bool:
    idx = find_idx(doc, anchor_text)
    if idx < 0:
        return False
    anchor = doc.paragraphs[idx]
    for line in lines:
        anchor.insert_paragraph_before(line)
    return True


def replace_block(doc: Document, start_text: str, end_text: str, lines: list[str]) -> bool:
    ok1 = remove_between(doc, start_text, end_text)
    ok2 = insert_lines_before(doc, end_text, lines)
    return ok1 and ok2


def replace_exact_line(doc: Document, old: str, new: str):
    for p in doc.paragraphs:
        if p_text(p) == old:
            p.text = new
            return True
    return False


def add_picture_before_heading(doc: Document, heading_text: str, image_path: Path, caption: str):
    idx = find_idx(doc, heading_text)
    if idx < 0:
        return False

    anchor = doc.paragraphs[idx]
    pic_para = anchor.insert_paragraph_before("")
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.2))
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cap_para = anchor.insert_paragraph_before(caption)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def add_screenshot_placeholders(doc: Document):
    replace_block(
        doc,
        "CHAPTER 4:  Input Screens and Reports",
        "CHAPTER 5:  Drawbacks and Limitations",
        [
            "4.1 :  INPUT FORMS WITH DATA",
            "(Insert your website input screenshots below)",
            "",
            "",
            "",
            "A) Login / Registration Input Screen",
            "",
            "",
            "",
            "B) Upload Form Input Screen",
            "",
            "",
            "",
            "4.2     OUTPUT REPORTS WITH DATA",
            "(Insert your website output screenshots below)",
            "",
            "",
            "",
            "A) Prediction Result Output Screen",
            "",
            "",
            "",
            "B) Case Detail / Patient Report Output Screen",
            "",
            "",
            "",
        ],
    )


def apply_content(doc: Document):
    replace_exact_line(doc, "SYNOPSIS REPORT", "PROJECT REPORT")
    replace_exact_line(doc, "Sakshi Pandey 1272240", "Sakshi Pandey 1272240631")

    replace_block(
        doc,
        "Project Report Index",
        "CHAPTER 1:  INTRODUCTION",
        [
            "CHAPTER 1:  INTRODUCTION",
            "1.1  Company Profile (only applicable for Live Project)",
            "1.2  Existing System and Need for System",
            "1.3  Scope of Work",
            "1.4  Operating Environment - Hardware and Software",
            "CHAPTER 2:  PROPOSED SYSTEM",
            "2.1 Proposed System",
            "2.2 Objectives of System",
            "CHAPTER 3:  ANALYSIS & DESIGN",
            "3.1 Entity Relationship Diagram",
            "3.2 Data Flow Diagrams",
            "CHAPTER 4:  Input Screens and Reports",
            "4.1 :  INPUT FORMS WITH DATA",
            "4.2     OUTPUT REPORTS WITH DATA",
            "CHAPTER 5:  Drawbacks and Limitations",
            "CHAPTER 6:  Proposed Enhancements",
            "CHAPTER 7:  Bibliography",
            "",
        ],
    )

    replace_block(
        doc,
        "CHAPTER 1:  INTRODUCTION",
        "CHAPTER 2:  PROPOSED SYSTEM",
        [
            "1.1  Company Profile (only applicable for Live Project)",
            "CuraVision AI is an academic healthcare-technology mini project that provides AI-assisted chest X-ray triage support through a web application. The solution integrates secure doctor workflow, machine-learning inference, persistent report storage, and patient report lookup.",
            "",
            "1.2  Existing System and Need for System",
            "Many educational medical-AI demos directly classify uploaded files without verifying whether the image is a valid chest X-ray. This can generate unreliable outputs and weakens trust in AI-assisted triage.",
            "CuraVision AI addresses this by using a two-stage pipeline: suitability validation first, pneumonia classification second. This improves practical reliability and creates a better clinician-facing workflow.",
            "",
            "1.3  Scope of Work",
            "The scope includes doctor registration and OTP verification, secure login, case upload with patient details, chest X-ray validation, prediction with confidence score, AI-generated summary, and persistent case/report retrieval.",
            "The project scope is educational and decision-support oriented; it does not provide final clinical diagnosis.",
            "",
            "1.4  Operating Environment - Hardware and Software",
            "Hardware: Minimum 8 GB RAM (16 GB recommended), modern multicore CPU, optional CUDA-enabled GPU for training experiments.",
            "Software: Python 3.11, Flask, Flask-SQLAlchemy, Flask-Bcrypt, TensorFlow/Keras, Transformers CLIP, Pillow/OpenCV, SQLite, Git, and VS Code.",
            "",
        ],
    )

    replace_block(
        doc,
        "CHAPTER 2:  PROPOSED SYSTEM",
        "CHAPTER 3:  ANALYSIS & DESIGN",
        [
            "2.1 Proposed System",
            "The proposed system is a web-based Clinical Decision Support System for chest X-ray triage. It performs model inference only after image suitability validation and presents transparent confidence-based outputs.",
            "The architecture includes authentication layer, upload and validation layer, AI inference layer, and persistence/reporting layer.",
            "",
            "2.2 Objectives of System",
            "1. Implement reliable two-stage inference (validation + classification).",
            "2. Reduce out-of-domain misuse by rejecting non-chest inputs.",
            "3. Provide confidence and explanatory summary for each prediction.",
            "4. Maintain secure doctor workflow with OTP verification.",
            "5. Store and retrieve complete case history for academic evaluation.",
            "",
        ],
    )

    replace_block(
        doc,
        "CHAPTER 3:  ANALYSIS & DESIGN",
        "CHAPTER 4:  Input Screens and Reports",
        [
            "3.1 Entity Relationship Diagram",
            "The database layer consists primarily of doctors and patients entities. Doctor records support authentication and verification lifecycle, while patient records store case metadata, prediction outputs, confidence values, and report timeline.",
            "",
            "3.2 Data Flow Diagrams",
            "Main flow: Doctor Login -> Upload Patient Data and X-ray -> Suitability Validation -> Classification -> Result + Confidence -> Save to Database -> Report Retrieval.",
            "Secondary flow: Patient ID -> Report Lookup -> Report Detail View.",
            "",
        ],
    )

    add_screenshot_placeholders(doc)

    replace_block(
        doc,
        "CHAPTER 5:  Drawbacks and Limitations",
        "CHAPTER 6:  Proposed Enhancements",
        [
            "The current system is limited to binary classification (NORMAL vs PNEUMONIA).",
            "Prediction quality depends on dataset representativeness and uploaded image quality.",
            "The deployment is academic and local; it is not a certified clinical system.",
            "Outputs require expert validation and should not be treated as final diagnosis.",
            "",
        ],
    )

    replace_block(
        doc,
        "CHAPTER 6:  Proposed Enhancements",
        "CHAPTER 7:  Bibliography",
        [
            "Extend to multi-class thoracic disease prediction.",
            "Add explainability visualization such as Grad-CAM.",
            "Introduce role-based access control and cloud deployment pipeline.",
            "Improve calibration and threshold tuning using validation cohorts.",
            "Integrate periodic retraining workflow with clinician feedback.",
            "",
        ],
    )

    # Bibliography content is appended by append_bibliography_if_needed().


def append_bibliography_if_needed(doc: Document):
    # If CHAPTER 7 is the last section, ensure bibliography lines exist even if replace_block with blank end failed.
    idx = find_idx(doc, "CHAPTER 7:  Bibliography")
    if idx < 0:
        return

    found_ref = False
    for p in doc.paragraphs[idx + 1 : min(len(doc.paragraphs), idx + 8)]:
        if p_text(p).startswith("1. Kaggle"):
            found_ref = True
            break

    if not found_ref:
        anchor = doc.paragraphs[idx + 1] if idx + 1 < len(doc.paragraphs) else doc.paragraphs[-1]
        refs = [
            "1. Kaggle Chest X-ray Pneumonia Dataset: https://www.kaggle.com/paultimothymooney/chest-xray-pneumonia",
            "2. TensorFlow Documentation: https://www.tensorflow.org/",
            "3. Keras Documentation: https://keras.io/",
            "4. Flask Documentation: https://flask.palletsprojects.com/",
            "5. Flask-SQLAlchemy Documentation: https://flask-sqlalchemy.palletsprojects.com/",
            "6. Transformers Documentation: https://huggingface.co/docs/transformers/",
            "7. MobileNetV2 Paper: https://arxiv.org/abs/1801.04381",
            "8. SQLite Documentation: https://www.sqlite.org/docs.html",
            "9. Project Repository: https://github.com/vedant420-s/mini-project",
        ]
        for line in refs:
            anchor.insert_paragraph_before(line)


def remove_all_inline_pictures(doc: Document):
    to_remove = []
    for p in doc.paragraphs:
        if p._element.xpath('.//w:drawing') or p._element.xpath('.//w:pict'):
            to_remove.append(p)
    for p in to_remove:
        remove_paragraph(p)


def main():
    for p in DIAGRAMS.values():
        if not p.exists():
            raise FileNotFoundError(f"Missing diagram: {p}")

    if not GUIDELINE_DOC.exists():
        raise FileNotFoundError(f"Guideline document not found: {GUIDELINE_DOC}")

    doc = Document(str(GUIDELINE_DOC))

    remove_all_inline_pictures(doc)
    apply_content(doc)
    append_bibliography_if_needed(doc)

    add_picture_before_heading(doc, "CHAPTER 4:  Input Screens and Reports", DIAGRAMS["architecture"], "Figure: System Architecture")
    add_picture_before_heading(doc, "CHAPTER 4:  Input Screens and Reports", DIAGRAMS["er"], "Figure: ER Diagram")
    add_picture_before_heading(doc, "CHAPTER 5:  Drawbacks and Limitations", DIAGRAMS["flow"], "Figure: Project Flow Diagram")
    add_picture_before_heading(doc, "CHAPTER 7:  Bibliography", DIAGRAMS["pipeline"], "Figure: AI Inference Pipeline")

    doc.save(str(OUT_DESKTOP))
    doc.save(str(OUT_WORKSPACE))

    final_doc = Document(str(OUT_DESKTOP))
    print(f"Saved (desktop): {OUT_DESKTOP}")
    print(f"Saved (workspace): {OUT_WORKSPACE}")
    print(f"Final image count: {len(final_doc.inline_shapes)}")


if __name__ == "__main__":
    main()
