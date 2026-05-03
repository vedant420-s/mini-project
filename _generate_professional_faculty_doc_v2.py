from __future__ import annotations

from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"c:\Users\VEDANT\OneDrive\Desktop\mini project")
DESKTOP = Path(r"c:\Users\VEDANT\OneDrive\Desktop")
ASSET_DIR = ROOT / "doc_diagrams" / "professional_faculty_v2"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DOC = DESKTOP / "AI_Medical_Image_Classifier_Faculty_Report.docx"

PROJECT_TITLE = "AI-Powered Clinical Decision Support System for Chest X-ray Triage"
TEAM = "Vedant and Sakshi"
REPO = "https://github.com/vedant420-s/mini-project"

PALETTE = {
    "navy": (23, 54, 93),
    "teal": (21, 125, 132),
    "orange": (211, 123, 53),
    "green": (62, 144, 91),
    "gray": (90, 98, 110),
    "line": (84, 102, 128),
    "bg1": (245, 248, 253),
    "bg2": (237, 248, 245),
    "bg3": (255, 245, 236),
}


def load_font(size: int, bold: bool = False):
    candidates = [
        "segoeuib.ttf" if bold else "segoeui.ttf",
        "cambria.ttf",
        "arialbd.ttf" if bold else "arial.ttf",
    ]
    for c in candidates:
        try:
            return ImageFont.truetype(c, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_wrapped(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], font, fill=(35, 35, 35), gap: int = 7):
    x1, y1, x2, y2 = box
    max_w = x2 - x1

    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_w:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    y = y1
    for ln in lines:
        draw.text((x1, y), ln, font=font, fill=fill)
        h = draw.textbbox((0, 0), ln, font=font)[3]
        y += h + gap
        if y > y2:
            break


def draw_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str, border, fill):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, outline=border, width=5, fill=fill)
    draw.text((x1 + 18, y1 + 14), title, font=load_font(31, bold=True), fill=border)
    draw_wrapped(draw, body, (x1 + 18, y1 + 60, x2 - 18, y2 - 16), load_font(23))


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(65, 85, 118), width: int = 6):
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill=color, width=width)
    h = 15
    if abs(ex - sx) > abs(ey - sy):
        if ex >= sx:
            pts = [(ex, ey), (ex - h, ey - h // 2), (ex - h, ey + h // 2)]
        else:
            pts = [(ex, ey), (ex + h, ey - h // 2), (ex + h, ey + h // 2)]
    else:
        if ey >= sy:
            pts = [(ex, ey), (ex - h // 2, ey - h), (ex + h // 2, ey - h)]
        else:
            pts = [(ex, ey), (ex - h // 2, ey + h), (ex + h // 2, ey + h)]
    draw.polygon(pts, fill=color)


def create_system_architecture(path: Path):
    img = Image.new("RGB", (2100, 1250), PALETTE["bg1"])
    draw = ImageDraw.Draw(img)

    draw.text((60, 24), "System Architecture", font=load_font(56, bold=True), fill=PALETTE["navy"])

    draw_box(
        draw,
        (90, 170, 650, 470),
        "Presentation Layer",
        "Doctor and patient interfaces built using Jinja templates, HTML, Tailwind utility classes, and JavaScript interactions.",
        PALETTE["navy"],
        (255, 255, 255),
    )
    draw_box(
        draw,
        (770, 170, 1330, 470),
        "Application Layer",
        "Flask blueprints process authentication, validation, uploads, reporting, and workflow navigation.",
        PALETTE["teal"],
        PALETTE["bg2"],
    )
    draw_box(
        draw,
        (1450, 170, 2010, 470),
        "AI Layer",
        "CLIP suitability detection + TensorFlow classifier + rule-based insight generation for explainable support.",
        PALETTE["orange"],
        PALETTE["bg3"],
    )

    draw_box(
        draw,
        (280, 650, 990, 1040),
        "Data Layer",
        "SQLite + Flask-SQLAlchemy stores doctor credentials, patient identifiers, image metadata, prediction confidence, and case timeline.",
        PALETTE["green"],
        (242, 252, 245),
    )
    draw_box(
        draw,
        (1120, 650, 1830, 1040),
        "Service Integration",
        "Email delivery for OTP and notification using Resend or SendGrid APIs with SMTP fallback.",
        PALETTE["gray"],
        (248, 250, 252),
    )

    draw_arrow(draw, (650, 320), (770, 320), PALETTE["line"], 6)
    draw_arrow(draw, (1330, 320), (1450, 320), PALETTE["line"], 6)
    draw_arrow(draw, (1050, 470), (700, 650), PALETTE["line"], 6)
    draw_arrow(draw, (1150, 470), (1460, 650), PALETTE["line"], 6)

    draw.text((60, 1170), "Figure: Layer-wise architecture used in the implemented system.", font=load_font(24), fill=PALETTE["gray"])
    img.save(path)


def create_workflow(path: Path):
    img = Image.new("RGB", (2100, 1320), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    draw.text((60, 24), "Clinical Workflow", font=load_font(56, bold=True), fill=PALETTE["navy"])

    steps = [
        ("1. Register and Verify", "Doctor creates account and verifies email through OTP."),
        ("2. Login and Access Dashboard", "Session-protected dashboard provides case statistics and alerts."),
        ("3. Upload Case", "Doctor submits patient details and chest X-ray file."),
        ("4. Suitability Gate", "CLIP confirms chest X-ray validity before diagnosis."),
        ("5. Predict and Explain", "TensorFlow classification + confidence + AI insight summary."),
        ("6. Persist and Review", "Case is saved and available in history, detail page, and patient portal."),
    ]

    y = 150
    for idx, (title, text) in enumerate(steps):
        top = y + idx * 180
        draw.rounded_rectangle((120, top, 1980, top + 130), radius=15, outline=PALETTE["teal"], width=4, fill=PALETTE["bg1"])
        draw.text((160, top + 16), title, font=load_font(30, bold=True), fill=PALETTE["navy"])
        draw.text((160, top + 72), text, font=load_font(24), fill=(40, 40, 40))
        if idx < len(steps) - 1:
            draw_arrow(draw, (1050, top + 130), (1050, top + 180), PALETTE["line"], 6)

    draw.rounded_rectangle((120, 1240, 1980, 1300), radius=12, outline=PALETTE["green"], width=3, fill=(239, 252, 243))
    draw.text((160, 1258), "Patient Portal Branch: Patient ID -> Reports -> Report Detail", font=load_font(24, bold=True), fill=PALETTE["green"])

    img.save(path)


def create_ai_pipeline(path: Path):
    img = Image.new("RGB", (2100, 1180), PALETTE["bg1"])
    draw = ImageDraw.Draw(img)

    draw.text((60, 24), "AI Inference Pipeline", font=load_font(56, bold=True), fill=PALETTE["navy"])

    blocks = [
        (120, "Input", "X-ray upload from doctor form"),
        (450, "Preprocess", "Resize and normalize image tensor"),
        (780, "CLIP Check", "Reject non-chest out-of-domain images"),
        (1110, "CNN Predict", "NORMAL/PNEUMONIA raw probability"),
        (1440, "Threshold", "Configured cutoff at 0.40"),
        (1770, "Output", "Confidence, message, and AI insights"),
    ]

    for x, title, text in blocks:
        draw.rounded_rectangle((x, 300, x + 250, 760), radius=16, outline=PALETTE["navy"], width=4, fill=(255, 255, 255))
        draw.text((x + 16, 350), title, font=load_font(29, bold=True), fill=PALETTE["navy"])
        draw_wrapped(draw, text, (x + 16, 430, x + 234, 730), load_font(22), fill=(45, 45, 45), gap=7)

    for i in range(len(blocks) - 1):
        x_start = blocks[i][0] + 250
        x_end = blocks[i + 1][0]
        draw_arrow(draw, (x_start, 530), (x_end, 530), PALETTE["line"], 6)

    draw.rounded_rectangle((120, 860, 2020, 1070), radius=16, outline=PALETTE["green"], width=4, fill=(239, 252, 243))
    draw_wrapped(
        draw,
        "After prediction, the system attaches generated symptoms and AI insight summary, commits record to SQLite, and returns JSON used in dashboard and case pages.",
        (160, 920, 1980, 1030),
        load_font(27),
        fill=PALETTE["green"],
        gap=9,
    )

    img.save(path)


def create_er_diagram(path: Path):
    img = Image.new("RGB", (2100, 1250), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    draw.text((60, 24), "ER Diagram and Database Model", font=load_font(56, bold=True), fill=PALETTE["navy"])

    draw.rounded_rectangle((100, 170, 930, 1090), radius=18, outline=PALETTE["navy"], width=5, fill=(241, 247, 255))
    draw.rounded_rectangle((1170, 170, 2000, 1090), radius=18, outline=PALETTE["teal"], width=5, fill=(237, 250, 244))

    draw.text((140, 210), "TABLE: doctors", font=load_font(38, bold=True), fill=PALETTE["navy"])
    draw.text((1210, 210), "TABLE: patients", font=load_font(38, bold=True), fill=PALETTE["teal"])

    left_lines = [
        "PK id: Integer",
        "name: String(120)",
        "email: String(255), unique, index",
        "password_hash: String(255)",
        "is_verified: Boolean",
        "verification_otp: String(10), nullable",
        "otp_expires_at: DateTime, nullable",
    ]

    right_lines = [
        "PK id: Integer",
        "patient_identifier: String(64), index",
        "name: String(120)",
        "age: Integer",
        "symptoms: Text",
        "generated_symptoms: Text",
        "image_path/photo_path: String(500)",
        "image_blob/photo_blob: LargeBinary",
        "prediction: String(50)",
        "confidence: Float",
        "ai_insights: Text",
        "created_at: DateTime",
    ]

    y = 290
    for ln in left_lines:
        draw.text((140, y), ln, font=load_font(27), fill=(35, 35, 35))
        y += 62

    y = 290
    for ln in right_lines:
        draw.text((1210, y), ln, font=load_font(27), fill=(35, 35, 35))
        y += 58

    draw_arrow(draw, (930, 620), (1170, 620), PALETTE["gray"], 7)
    draw.text((850, 680), "application workflow linkage", font=load_font(24), fill=PALETTE["gray"])

    draw.text((60, 1170), "SQLite with Flask-SQLAlchemy ORM for persistent case and authentication management.", font=load_font(24), fill=PALETTE["gray"])
    img.save(path)


def set_page_border(section):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")

    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "10")
        elem.set(qn("w:space"), "24")
        elem.set(qn("w:color"), "264478")
        pg_borders.append(elem)

    sect_pr.append(pg_borders)


def color_heading(paragraph, rgb: tuple[int, int, int]):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = "Cambria"
        run.font.bold = True
        run.font.color.rgb = RGBColor(*rgb)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.italic = True
    r.font.color.rgb = RGBColor(*PALETTE["gray"])


def add_para(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25


def add_bullet_lines(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(line, style="List Bullet")
        p.paragraph_format.space_after = Pt(5)


def add_number_lines(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph(line, style="List Number")
        p.paragraph_format.space_after = Pt(5)


def long_paragraphs() -> dict[str, list[str]]:
    return {
        "abstract": [
            "This report documents the design and implementation of an educational Clinical Decision Support System that assists in chest X-ray triage using a two-stage inference pipeline. The system combines a lightweight web interface, secure doctor authentication, data persistence, and machine learning inference into a single integrated workflow intended for academic demonstration.",
            "In contrast to basic student prototypes that directly classify every uploaded image, the implemented solution introduces an explicit suitability gate before prediction. This design decision improves reliability by detecting irrelevant images and allowing only chest X-rays to proceed to pneumonia classification. The approach demonstrates not just model usage but practical system-level safeguards.",
            "The project also emphasizes usability and traceability. Every prediction is tied to patient details, manual symptoms, generated symptom hints, confidence score, and timestamp. Cases can be reviewed through dashboard and detail pages, and patients can access reports with identifier-based lookup. As a result, the system behaves like a complete workflow application instead of an isolated model demo.",
        ],
        "introduction": [
            "Artificial intelligence is increasingly used to support healthcare decision processes, especially in imaging where interpretation speed and consistency are important. Chest X-ray analysis is a frequent use case because pneumonia screening is clinically relevant and datasets are available for academic research. However, trustworthy adoption requires a clear separation between assistance and diagnosis.",
            "The present mini project is built with that principle in mind. The system provides preliminary risk indications but explicitly avoids replacing professional judgment. The application includes a visible medical disclaimer and encourages clinician confirmation for all outcomes. This positioning is essential for responsible AI use in educational settings.",
            "From a software engineering perspective, the project integrates front-end interaction, backend business logic, model serving, persistence, and notifications. Such full-stack integration is useful for evaluation because it demonstrates architecture design, data flow management, validation controls, and maintainable coding structure under one unified objective.",
        ],
        "problem": [
            "Many beginner medical-AI projects suffer from three major limitations: weak input validation, absence of secure user workflows, and poor reporting continuity. These issues lead to unrealistic demonstrations where any random image receives a medical label, user identity is not protected, and outputs are not stored in a reproducible format.",
            "The current project addresses this gap by structuring the pipeline as: authenticated access, validated upload, suitability check, prediction with confidence, explanatory summary, and persistent storage. This sequence mirrors practical workflow requirements and reduces avoidable misuse.",
            "Another challenge is balancing sensitivity and specificity for educational interpretation. The model threshold and confidence display are designed to make output behavior transparent. Instead of hiding uncertainty, the application exposes score values and contextual messaging to help interpret results responsibly.",
        ],
        "requirements": [
            "Functional requirements include doctor registration and login, OTP email verification, case creation with image upload, prediction generation, case listing, case-detail retrieval, and patient report access using patient identifiers. The system must also support deletion of case records and secure retrieval of stored images.",
            "Non-functional requirements include usability, response clarity, maintainability, and controlled failure behavior. Validation errors should be human readable; invalid file types should be rejected early; and exception paths should preserve database consistency via rollback routines.",
            "Operational requirements include Python runtime compatibility, dependency installation through requirements.txt, model artifact availability, and local database initialization at startup. The application should run reliably in a standard Windows developer environment with minimal setup steps.",
        ],
        "dataset": [
            "The project references the publicly available chest X-ray pneumonia dataset used widely in educational research. The data is inherently imbalanced, with pneumonia instances outnumbering normal class samples. Such imbalance influences precision and recall trade-offs and must be acknowledged during interpretation.",
            "Preprocessing aligns images to the model input shape and normalizes pixel values. Controlled preprocessing ensures consistency across training and inference. At application level, uploaded files are further validated by extension checks and a semantic suitability gate before reaching the classifier.",
            "Dataset characteristics motivate careful result interpretation. High recall for pneumonia can be clinically useful in triage scenarios, while false positives remain acceptable only when downstream human review is guaranteed. The report therefore treats model outputs as decision-support indicators, not definitive diagnoses.",
        ],
        "methodology": [
            "The adopted methodology follows a layered and modular pattern. Authentication responsibilities are separated into dedicated routes, clinical workflows are concentrated in the main route module, persistence is abstracted through SQLAlchemy models, and helper logic is isolated in AI utility modules.",
            "Inference itself is implemented in two stages. The first stage verifies whether an uploaded image resembles a chest X-ray using CLIP text-image similarity. Only if this stage succeeds does the second stage run TensorFlow-based binary classification for NORMAL versus PNEUMONIA.",
            "Post-inference, the system computes confidence, generates symptom-oriented insight text, persists all relevant case fields, and returns a structured response for UI rendering. This chain demonstrates a complete and explainable processing lifecycle from input to recordable output.",
        ],
        "modules": [
            "The authentication module manages registration, OTP generation, verification expiry logic, login, and session state. Route protection is implemented using a reusable decorator that prevents unauthorized access to sensitive pages.",
            "The main routes module orchestrates upload handling, preprocessing, suitability detection, prediction, response construction, and history retrieval. It also serves case images from database blobs with fallback to file storage, ensuring robust media access behavior.",
            "The model module defines Doctor and Patient entities with explicit field typing, constraints, and defaults. This schema provides consistency for credentials, report metadata, image storage, confidence tracking, and timestamped audit history.",
        ],
        "testing": [
            "Testing combines script-level checks and manual workflow validation. Automated tests verify endpoint behavior and key pipeline branches. Manual tests validate user journeys such as registration, verification, upload, prediction visibility, and case review navigation.",
            "Error path testing is equally important. Invalid age formats, missing symptoms, unsupported file types, or non-chest images must trigger informative messages without corrupting state. Such tests ensure the application fails safely and predictably.",
            "Performance evaluation uses reported accuracy and class-wise metrics from project artifacts. The report interprets these metrics with medical caution, emphasizing that high recall supports triage sensitivity while final diagnostic responsibility remains with clinicians.",
        ],
        "security": [
            "Security measures include session-based authorization for protected routes, password hashing, OTP-based verification, controlled file handling through secure filenames, and strict patient identifier validation patterns.",
            "Data reliability is supported by commit and rollback semantics in database transactions. In case of runtime exceptions during case creation, the application returns a controlled error response and avoids partial writes.",
            "Although this is an educational deployment, the architecture demonstrates core secure-coding practices expected in real systems: input sanitization, authentication checks, principle of least privilege in route exposure, and transparent error handling.",
        ],
        "limitations": [
            "The project currently focuses on binary classification only, which limits disease coverage. Clinical images can present complex comorbidities and artifact conditions that require broader multi-class or multi-label modeling.",
            "Inference quality depends on dataset representativeness and acquisition variability. Domain shifts, poor image quality, and equipment differences can reduce performance when moving beyond controlled datasets.",
            "The deployment is local and academic; therefore, advanced production concerns such as distributed scaling, enterprise authentication, centralized observability, and compliance frameworks are outside current scope.",
        ],
        "future": [
            "Future enhancements include expanding to multi-class thoracic findings, integrating explainability maps such as Grad-CAM, and calibrating confidence thresholds against validation cohorts for improved interpretability.",
            "A cloud deployment path can introduce containerization, role-based access controls, and secure object storage for media assets. This would improve portability, reliability, and operational governance.",
            "Another valuable extension is active-learning feedback loops where clinician corrections are captured to guide periodic retraining and improve model robustness over time.",
        ],
        "conclusion": [
            "This mini project demonstrates a complete academic AI application that goes beyond simple model inference. It unifies secure access, controlled input validation, staged prediction logic, persistent case management, and user-facing output interpretation.",
            "By combining engineering structure with responsible-AI messaging, the implementation offers strong educational value for final evaluation. It shows practical understanding of both machine-learning integration and healthcare-oriented software design constraints.",
            "With planned improvements in explainability, deployment, and dataset diversity, the platform can evolve into a stronger research prototype while preserving its clinician-in-the-loop design philosophy.",
        ],
    }


def generate_diagrams() -> dict[str, Path]:
    paths = {
        "architecture": ASSET_DIR / "system_architecture.png",
        "workflow": ASSET_DIR / "clinical_workflow.png",
        "pipeline": ASSET_DIR / "ai_pipeline.png",
        "er": ASSET_DIR / "er_diagram.png",
    }
    create_system_architecture(paths["architecture"])
    create_workflow(paths["workflow"])
    create_ai_pipeline(paths["pipeline"])
    create_er_diagram(paths["er"])
    return paths


def add_screenshot_space_page(doc: Document, heading_text: str, figure_title: str, figure_no: str):
    h = doc.add_heading(heading_text, level=2)
    color_heading(h, PALETTE["navy"])

    note = doc.add_paragraph(
        "Reserved page for final project screenshot. Replace this space with your original working output before submission."
    )
    note.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    row = table.rows[0]
    row.height = Inches(5.7)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    cell.text = "\n\n\n\n\n\n\n\n\n\n\n\nSCREENSHOT PLACEHOLDER\n\n(Insert project screen capture here)"
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = "Calibri"
            r.font.size = Pt(13)
            r.font.bold = True
            r.font.color.rgb = RGBColor(*PALETTE["gray"])

    add_caption(doc, f"Figure {figure_no}: {figure_title}")


def build_document(diagrams: dict[str, Path]):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    set_page_border(section)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    title = doc.add_paragraph("AI MEDICAL IMAGE CLASSIFIER")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.name = "Cambria"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*PALETTE["navy"])

    subtitle = doc.add_paragraph("Detailed Project Report for Faculty Evaluation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    srun = subtitle.runs[0]
    srun.font.name = "Cambria"
    srun.font.size = Pt(16)
    srun.font.italic = True
    srun.font.color.rgb = RGBColor(*PALETTE["teal"])

    doc.add_paragraph("")

    info = doc.add_table(rows=6, cols=2)
    info.style = "Table Grid"
    info.cell(0, 0).text = "Project Title"
    info.cell(0, 1).text = PROJECT_TITLE
    info.cell(1, 0).text = "Team"
    info.cell(1, 1).text = TEAM
    info.cell(2, 0).text = "Repository"
    info.cell(2, 1).text = REPO
    info.cell(3, 0).text = "Course Component"
    info.cell(3, 1).text = "Mini Project / PBL"
    info.cell(4, 0).text = "Submission Date"
    info.cell(4, 1).text = datetime.now().strftime("%d %B %Y")
    info.cell(5, 0).text = "Document Version"
    info.cell(5, 1).text = "Faculty Report - Extended Edition"

    doc.add_page_break()

    h = doc.add_heading("Certificate of Original Work", level=1)
    color_heading(h, PALETTE["navy"])
    add_para(
        doc,
        "We hereby declare that this report and accompanying implementation titled 'AI-Powered Clinical Decision Support System for Chest X-ray Triage' are submitted as part of the mini project requirement. The work presented is implemented and integrated by the project team for academic evaluation."
    )
    add_para(
        doc,
        "The project is intended for educational use only and not for real-world diagnostic deployment. All clinical outputs are advisory and require interpretation by qualified medical professionals."
    )

    sign_table = doc.add_table(rows=2, cols=2)
    sign_table.style = "Table Grid"
    sign_table.cell(0, 0).text = "Student Signature"
    sign_table.cell(0, 1).text = "Guide Signature"
    sign_table.cell(1, 0).text = "\n\n"
    sign_table.cell(1, 1).text = "\n\n"

    doc.add_page_break()

    h = doc.add_heading("Acknowledgement", level=1)
    color_heading(h, PALETTE["navy"])
    add_para(
        doc,
        "We express sincere gratitude to our faculty mentors and department for guidance throughout this project. Their support in problem framing, model evaluation, and software documentation significantly improved the quality and clarity of this work."
    )
    add_para(
        doc,
        "We also acknowledge the open-source ecosystem, including Python, Flask, TensorFlow, and public research datasets, which enabled rapid experimentation and implementation of this integrated educational solution."
    )

    doc.add_page_break()

    h = doc.add_heading("Table of Contents", level=1)
    color_heading(h, PALETTE["navy"])
    toc_lines = [
        "1. Abstract",
        "2. Introduction",
        "3. Problem Statement and Objectives",
        "4. Requirement Analysis",
        "5. Dataset and Preprocessing",
        "6. Methodology and System Design",
        "7. System Architecture Diagram",
        "8. Functional Module Explanation",
        "9. AI Inference Pipeline Diagram",
        "10. Workflow Diagram",
        "11. Database Design and ER Diagram",
        "12. Implementation Details",
        "13. Testing Strategy and Evaluation",
        "14. Security and Reliability",
        "15. Deployment and Execution Guide",
        "16. Limitations",
        "17. Future Scope",
        "18. Conclusion",
        "19. References",
        "20. Appendix: Input/Output Screenshots",
    ]
    add_number_lines(doc, toc_lines)

    text = long_paragraphs()

    h = doc.add_heading("1. Abstract", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["abstract"]:
        add_para(doc, p)

    h = doc.add_heading("2. Introduction", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["introduction"]:
        add_para(doc, p)

    h = doc.add_heading("3. Problem Statement and Objectives", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["problem"]:
        add_para(doc, p)
    add_bullet_lines(
        doc,
        [
            "Design a complete clinical decision-support web workflow.",
            "Ensure robust input validation before model inference.",
            "Provide transparent confidence-based predictions.",
            "Persist case history for longitudinal review.",
            "Maintain academic ethical boundaries via clear disclaimer.",
        ],
    )

    h = doc.add_heading("4. Requirement Analysis", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["requirements"]:
        add_para(doc, p)

    req_table = doc.add_table(rows=7, cols=3)
    req_table.style = "Table Grid"
    req_table.cell(0, 0).text = "ID"
    req_table.cell(0, 1).text = "Requirement"
    req_table.cell(0, 2).text = "Type"
    entries = [
        ("R1", "Doctor registration and login", "Functional"),
        ("R2", "OTP verification workflow", "Functional"),
        ("R3", "Two-stage X-ray inference", "Functional"),
        ("R4", "Case history persistence", "Functional"),
        ("R5", "Secure route access", "Non-Functional"),
        ("R6", "Readable error handling", "Non-Functional"),
    ]
    for i, (rid, req, typ) in enumerate(entries, start=1):
        req_table.cell(i, 0).text = rid
        req_table.cell(i, 1).text = req
        req_table.cell(i, 2).text = typ

    h = doc.add_heading("5. Dataset and Preprocessing", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["dataset"]:
        add_para(doc, p)
    add_bullet_lines(
        doc,
        [
            "Dataset source: Kaggle Chest X-ray Pneumonia Dataset.",
            "Image normalization to 0-1 scale for model compatibility.",
            "Input resizing to 224 x 224 pixels.",
            "Class imbalance considered during interpretation.",
        ],
    )

    h = doc.add_heading("6. Methodology and System Design", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["methodology"]:
        add_para(doc, p)

    h = doc.add_heading("7. System Architecture Diagram", level=1)
    color_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["architecture"]), width=Inches(6.8))
    add_caption(doc, "Figure 7.1: Layer-wise system architecture")
    add_para(
        doc,
        "The architecture separates interface, application logic, inference, persistence, and integration services to keep responsibilities modular and maintainable. This separation also simplifies testing and future enhancement planning."
    )

    h = doc.add_heading("8. Functional Module Explanation", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["modules"]:
        add_para(doc, p)

    module_table = doc.add_table(rows=6, cols=2)
    module_table.style = "Table Grid"
    module_table.cell(0, 0).text = "Module"
    module_table.cell(0, 1).text = "Responsibility"
    module_table.cell(1, 0).text = "auth.py"
    module_table.cell(1, 1).text = "Registration, OTP, login/logout, session creation"
    module_table.cell(2, 0).text = "routes.py"
    module_table.cell(2, 1).text = "Upload, prediction, case workflows, patient portal"
    module_table.cell(3, 0).text = "models.py"
    module_table.cell(3, 1).text = "Doctor and patient SQLAlchemy schema"
    module_table.cell(4, 0).text = "ai_logic.py"
    module_table.cell(4, 1).text = "AI helper assessment and insight generation"
    module_table.cell(5, 0).text = "email_utils.py"
    module_table.cell(5, 1).text = "OTP and notification email dispatch"

    h = doc.add_heading("9. AI Inference Pipeline Diagram", level=1)
    color_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["pipeline"]), width=Inches(6.8))
    add_caption(doc, "Figure 9.1: Two-stage inference pipeline")
    add_para(
        doc,
        "The first-stage gate reduces accidental misuse and improves trustworthiness of downstream predictions. The second-stage classifier applies threshold logic and exposes confidence values, helping users understand prediction certainty and decision context."
    )

    h = doc.add_heading("10. Workflow Diagram", level=1)
    color_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["workflow"]), width=Inches(6.8))
    add_caption(doc, "Figure 10.1: End-to-end doctor and patient report workflow")
    add_para(
        doc,
        "The workflow highlights complete lifecycle continuity from registration to persistent records. This continuity is a key reason the project is suited for faculty evaluation, since it demonstrates practical software flow and not just a model endpoint."
    )

    h = doc.add_heading("11. Database Design and ER Diagram", level=1)
    color_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["er"]), width=Inches(6.8))
    add_caption(doc, "Figure 11.1: ER diagram of doctors and patients models")
    add_para(
        doc,
        "The doctors table supports secure account lifecycle, while the patients table is designed for detailed report persistence including media and explanation fields. This schema enables traceability and case-wise historical analysis."
    )

    h = doc.add_heading("12. Implementation Details", level=1)
    color_heading(h, PALETTE["navy"])
    impl_lines = [
        "Input validation includes patient name, numeric age boundaries, symptom presence, and controlled identifier format.",
        "File handling uses secure filename generation with timestamped storage paths.",
        "Prediction response returns class, confidence, threshold, generated symptom text, and timestamp.",
        "Image serving supports blob-first retrieval with filesystem fallback.",
        "Case deletion routine removes both record and associated stored files safely.",
    ]
    for line in impl_lines:
        add_para(doc, line)

    route_table = doc.add_table(rows=10, cols=2)
    route_table.style = "Table Grid"
    route_table.cell(0, 0).text = "Route"
    route_table.cell(0, 1).text = "Purpose"
    route_rows = [
        ("/auth/register", "Create doctor account"),
        ("/auth/verify-email", "Validate OTP"),
        ("/auth/login", "Start authenticated session"),
        ("/dashboard", "Display analytics and recent cases"),
        ("/upload", "Upload case form"),
        ("/predict (POST)", "Inference and persistence"),
        ("/cases", "List all cases"),
        ("/cases/<id>", "Detailed case view"),
        ("/patient/reports/<patient_id>", "Patient report listing"),
    ]
    for i, (route, purpose) in enumerate(route_rows, start=1):
        route_table.cell(i, 0).text = route
        route_table.cell(i, 1).text = purpose

    h = doc.add_heading("13. Testing Strategy and Evaluation", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["testing"]:
        add_para(doc, p)

    metric_table = doc.add_table(rows=6, cols=2)
    metric_table.style = "Table Grid"
    metric_table.cell(0, 0).text = "Evaluation Metric"
    metric_table.cell(0, 1).text = "Value"
    metrics = [
        ("Training Accuracy", "92.64%"),
        ("Validation Accuracy", "87.66%"),
        ("Test Accuracy", "87.66%"),
        ("NORMAL F1", "81.53%"),
        ("PNEUMONIA F1", "90.73%"),
    ]
    for i, (k, v) in enumerate(metrics, start=1):
        metric_table.cell(i, 0).text = k
        metric_table.cell(i, 1).text = v

    h = doc.add_heading("14. Security and Reliability", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["security"]:
        add_para(doc, p)
    add_bullet_lines(
        doc,
        [
            "OTP expiry policy and resend mechanism.",
            "Authenticated-only access to clinical pages.",
            "Input sanitization and strict format checks.",
            "Exception handling with safe rollback behavior.",
        ],
    )

    h = doc.add_heading("15. Deployment and Execution Guide", level=1)
    color_heading(h, PALETTE["navy"])
    add_para(doc, "The project is executed in a Python environment with dependency installation through requirements.txt and model artifacts loaded from local models directory.")

    deploy_steps = [
        "Create and activate virtual environment.",
        "Install dependencies using pip install -r requirements.txt.",
        "Configure environment variables for secret key and mail provider.",
        "Ensure model artifacts are available in models directory.",
        "Run application using python app.py.",
        "Validate endpoints through login and upload workflows.",
    ]
    add_number_lines(doc, deploy_steps)

    h = doc.add_heading("16. Limitations", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["limitations"]:
        add_para(doc, p)

    h = doc.add_heading("17. Future Scope", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["future"]:
        add_para(doc, p)

    h = doc.add_heading("18. Conclusion", level=1)
    color_heading(h, PALETTE["navy"])
    for p in text["conclusion"]:
        add_para(doc, p)

    h = doc.add_heading("19. References", level=1)
    color_heading(h, PALETTE["navy"])
    refs = [
        "Flask Documentation: https://flask.palletsprojects.com/",
        "TensorFlow Documentation: https://www.tensorflow.org/",
        "Transformers Documentation: https://huggingface.co/docs/transformers/",
        "Kaggle Chest X-ray Pneumonia Dataset: https://www.kaggle.com/paultimothymooney/chest-xray-pneumonia",
        "Project Repository: https://github.com/vedant420-s/mini-project",
    ]
    add_bullet_lines(doc, refs)

    h = doc.add_heading("20. Appendix: Input/Output Screenshots", level=1)
    color_heading(h, PALETTE["navy"])
    add_para(doc, "The next pages are reserved for your final working screenshots. These should be replaced with actual project run captures before submission.")

    add_screenshot_space_page(doc, "20.1 Input Screen - Login / Upload", "Input screen placeholder", "20.1")
    doc.add_page_break()
    add_screenshot_space_page(doc, "20.2 Output Screen - Prediction / Dashboard", "Output screen placeholder", "20.2")

    doc.save(str(OUTPUT_DOC))


def main():
    diagrams = generate_diagrams()
    build_document(diagrams)
    print(f"Generated: {OUTPUT_DOC}")
    print(f"Diagrams generated (4 total) in: {ASSET_DIR}")


if __name__ == "__main__":
    main()
