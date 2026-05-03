from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


ROOT = Path(r"c:\Users\VEDANT\OneDrive\Desktop\mini project")
DIAGRAM_DIR = ROOT / "doc_diagrams"
DIAGRAM_DIR.mkdir(exist_ok=True)

DB_DIAGRAM = DIAGRAM_DIR / "database_design_diagram.png"
FLOW_DIAGRAM = DIAGRAM_DIR / "project_flow_diagram.png"

DOC_FILES = [
    ROOT / "Mini Project.docx",
    ROOT / "miniproject.docx",
]


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def load_font(size: int):
    candidates = [
        "arial.ttf",
        "calibri.ttf",
        "segoeui.ttf",
    ]
    for name in candidates:
        try:
            return ImageFont.truetype(name, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_wrapped_text(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], font, fill=(20, 20, 20), line_gap: int = 8):
    x1, y1, x2, y2 = box
    max_width = x2 - x1

    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        trial = word if not current else f"{current} {word}"
        w = draw.textbbox((0, 0), trial, font=font)[2]
        if w <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    y = y1
    for line in lines:
        draw.text((x1, y), line, font=font, fill=fill)
        line_h = draw.textbbox((0, 0), line, font=font)[3]
        y += line_h + line_gap
        if y > y2:
            break


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(40, 70, 120), width: int = 6):
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill=color, width=width)

    # Arrow head
    head = 14
    if abs(ex - sx) > abs(ey - sy):
        # Horizontal arrow
        if ex >= sx:
            pts = [(ex, ey), (ex - head, ey - head // 2), (ex - head, ey + head // 2)]
        else:
            pts = [(ex, ey), (ex + head, ey - head // 2), (ex + head, ey + head // 2)]
    else:
        # Vertical arrow
        if ey >= sy:
            pts = [(ex, ey), (ex - head // 2, ey - head), (ex + head // 2, ey - head)]
        else:
            pts = [(ex, ey), (ex - head // 2, ey + head), (ex + head // 2, ey + head)]
    draw.polygon(pts, fill=color)


def create_database_diagram(path: Path):
    img = Image.new("RGB", (1800, 1100), "white")
    draw = ImageDraw.Draw(img)

    title_font = load_font(54)
    heading_font = load_font(34)
    text_font = load_font(24)
    small_font = load_font(22)

    draw.text((80, 40), "Database Design (ER Overview)", font=title_font, fill=(10, 60, 120))

    left_box = (80, 180, 850, 940)
    right_box = (950, 180, 1720, 940)

    draw.rounded_rectangle(left_box, radius=24, outline=(10, 90, 170), width=5, fill=(239, 248, 255))
    draw.rounded_rectangle(right_box, radius=24, outline=(20, 120, 80), width=5, fill=(239, 255, 246))

    draw.text((120, 220), "TABLE: doctors", font=heading_font, fill=(10, 90, 170))
    draw.text((990, 220), "TABLE: patients", font=heading_font, fill=(20, 120, 80))

    left_text = (
        "PK  id : Integer\n"
        "name : String(120)\n"
        "email : String(255), UNIQUE, INDEX\n"
        "password_hash : String(255)\n"
        "is_verified : Boolean\n"
        "verification_otp : String(10), nullable\n"
        "otp_expires_at : DateTime, nullable"
    )
    right_text = (
        "PK  id : Integer\n"
        "patient_identifier : String(64), INDEX\n"
        "name : String(120)\n"
        "age : Integer\n"
        "symptoms : Text\n"
        "generated_symptoms : Text\n"
        "image_path, photo_path : String(500)\n"
        "image_blob, photo_blob : LargeBinary\n"
        "image_mime, photo_mime : String(100)\n"
        "prediction : String(50)\n"
        "confidence : Float\n"
        "ai_insights : Text\n"
        "created_at : DateTime"
    )

    draw.multiline_text((120, 290), left_text, font=text_font, fill=(30, 30, 30), spacing=10)
    draw.multiline_text((990, 290), right_text, font=text_font, fill=(30, 30, 30), spacing=10)

    # Conceptual relationship note (no direct foreign key in current schema)
    draw_arrow(draw, (850, 560), (950, 560), color=(110, 110, 110), width=5)
    draw.text((740, 600), "Application workflow link", font=small_font, fill=(70, 70, 70))
    draw.text((700, 635), "(doctor session creates patient case records)", font=small_font, fill=(70, 70, 70))

    footer = "Storage engine: SQLite   |   ORM: Flask-SQLAlchemy   |   Schema migration via startup ALTER checks"
    draw.text((80, 1010), footer, font=small_font, fill=(60, 60, 60))

    img.save(path)


def create_flow_diagram(path: Path):
    img = Image.new("RGB", (1900, 1250), "white")
    draw = ImageDraw.Draw(img)

    title_font = load_font(54)
    box_title_font = load_font(27)
    box_text_font = load_font(21)

    draw.text((80, 30), "Project Flow Diagram", font=title_font, fill=(90, 30, 140))

    def box(x1, y1, x2, y2, title, body, color):
        draw.rounded_rectangle((x1, y1, x2, y2), radius=22, outline=color, width=5, fill=(250, 248, 255))
        draw.text((x1 + 20, y1 + 14), title, font=box_title_font, fill=color)
        draw_wrapped_text(draw, body, (x1 + 20, y1 + 62, x2 - 20, y2 - 16), box_text_font)

    box(100, 130, 560, 300, "1) Doctor Access", "Register, verify OTP, and login to open dashboard workflow.", (95, 52, 170))
    box(100, 370, 560, 540, "2) Case Upload", "Enter patient details and upload chest X-ray (plus optional photo).", (95, 52, 170))
    box(100, 610, 560, 780, "3) Stage-1 Validation", "CLIP suitability check rejects non-chest images (OOD control).", (95, 52, 170))
    box(100, 850, 560, 1020, "4) Stage-2 Prediction", "TensorFlow model predicts NORMAL/PNEUMONIA + confidence.", (95, 52, 170))

    box(700, 250, 1180, 430, "5) AI Insights", "Rule-based helper generates supportive clinical text and guidance.", (33, 120, 100))
    box(1320, 250, 1800, 430, "6) Persistence", "Save case record, metadata, images, and prediction in SQLite.", (33, 120, 100))
    box(700, 620, 1180, 800, "7) Doctor Review", "Dashboard, case history, high-risk list, and case detail screens.", (200, 110, 20))
    box(1320, 620, 1800, 800, "8) Patient Portal", "Patient ID based report lookup and detail view.", (200, 110, 20))

    draw_arrow(draw, (330, 300), (330, 370))
    draw_arrow(draw, (330, 540), (330, 610))
    draw_arrow(draw, (330, 780), (330, 850))
    draw_arrow(draw, (560, 930), (700, 340))
    draw_arrow(draw, (1180, 340), (1320, 340))
    draw_arrow(draw, (1560, 430), (1560, 620))
    draw_arrow(draw, (1320, 710), (1180, 710))

    draw.text((700, 980), "Outcome: Clinician-facing prediction, confidence, stored reports, and reviewable workflow", font=load_font(24), fill=(60, 60, 60))

    img.save(path)


def iter_paragraphs(document: Document):
    for p in document.paragraphs:
        yield p


def paragraph_text(p) -> str:
    return (p.text or "").strip()


def find_paragraph_index(document: Document, exact_text: str) -> int:
    for idx, p in enumerate(document.paragraphs):
        if paragraph_text(p) == exact_text:
            return idx
    return -1


def remove_between(document: Document, start_idx: int, end_idx: int):
    # Remove paragraphs strictly between start and end.
    to_remove = document.paragraphs[start_idx + 1 : end_idx]
    for p in to_remove:
        el = p._element
        el.getparent().remove(el)


def insert_paragraph_before(anchor_para, text: str = "", align_left: bool = True):
    p = anchor_para.insert_paragraph_before(text)
    if align_left:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return p


def insert_image_before(anchor_para, image_path: Path, width_inches: float = 6.5):
    p = anchor_para.insert_paragraph_before("")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def update_docx(doc_path: Path):
    doc = Document(str(doc_path))

    idx_db_design = find_paragraph_index(doc, "6.4 Database Design")
    idx_db_tables = find_paragraph_index(doc, "6.5 Database Tables")
    idx_flow = find_paragraph_index(doc, "6.6 Project Flow Diagram")
    idx_next = find_paragraph_index(doc, "6.7 Testing and Deployment")

    if min(idx_db_design, idx_db_tables, idx_flow, idx_next) < 0:
        raise RuntimeError(f"Required headings not found in {doc_path.name}")

    # Refresh indices after first removal.
    remove_between(doc, idx_db_design, idx_db_tables)

    # Recompute to avoid stale indices.
    idx_db_tables = find_paragraph_index(doc, "6.5 Database Tables")
    anchor_db_tables = doc.paragraphs[idx_db_tables]
    insert_paragraph_before(anchor_db_tables, "Figure 6.4: Database Design Diagram")
    insert_image_before(anchor_db_tables, DB_DIAGRAM, width_inches=7.2)
    insert_paragraph_before(anchor_db_tables, "The ER-style diagram above shows key entities, important attributes, and the workflow relationship used by the current application design.")

    idx_flow = find_paragraph_index(doc, "6.6 Project Flow Diagram")
    idx_next = find_paragraph_index(doc, "6.7 Testing and Deployment")
    remove_between(doc, idx_flow, idx_next)

    idx_next = find_paragraph_index(doc, "6.7 Testing and Deployment")
    anchor_next = doc.paragraphs[idx_next]
    insert_paragraph_before(anchor_next, "Figure 6.6: End-to-End Project Flow")
    insert_image_before(anchor_next, FLOW_DIAGRAM, width_inches=7.3)
    insert_paragraph_before(anchor_next, "The flow diagram illustrates doctor onboarding, image validation, prediction, insight generation, storage, and report access steps.")

    try:
        doc.save(str(doc_path))
        return doc_path
    except PermissionError:
        fallback = doc_path.with_name(f"{doc_path.stem}.with_diagrams{doc_path.suffix}")
        doc.save(str(fallback))
        return fallback


if __name__ == "__main__":
    create_database_diagram(DB_DIAGRAM)
    create_flow_diagram(FLOW_DIAGRAM)

    for doc_file in DOC_FILES:
        if doc_file.exists():
            output = update_docx(doc_file)
            print(f"Updated diagrams in: {output}")
        else:
            print(f"Skipped missing file: {doc_file}")

    print(f"Database diagram: {DB_DIAGRAM}")
    print(f"Flow diagram: {FLOW_DIAGRAM}")
