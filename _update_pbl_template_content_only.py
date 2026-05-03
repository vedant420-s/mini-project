from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(r"c:\Users\VEDANT\OneDrive\Desktop\mini project")
SOURCE_DOC = ROOT / "PBL.docx"
OUTPUT_WORKSPACE = ROOT / "PBL.docx"
OUTPUT_DESKTOP = Path(r"c:\Users\VEDANT\OneDrive\Desktop\PBL.docx")

DIAGRAM_DIR = ROOT / "doc_diagrams" / "professional_faculty_v2"
DIAGRAMS = {
    "system": DIAGRAM_DIR / "system_architecture.png",
    "er": DIAGRAM_DIR / "er_diagram.png",
    "flow": DIAGRAM_DIR / "clinical_workflow.png",
    "pipeline": DIAGRAM_DIR / "ai_pipeline.png",
}


def p_text(p) -> str:
    return (p.text or "").strip()


def find_idx(doc: Document, exact_text: str, start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        if p_text(doc.paragraphs[i]) == exact_text:
            return i
    return -1


def find_any_idx(doc: Document, candidates: list[str], start: int = 0) -> int:
    for i in range(start, len(doc.paragraphs)):
        current = p_text(doc.paragraphs[i])
        if current in candidates:
            return i
    return -1


def remove_paragraph(paragraph):
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_between(doc: Document, start_heading: str, end_heading_candidates: list[str]):
    s = find_idx(doc, start_heading)
    if s < 0:
        return False

    e = find_any_idx(doc, end_heading_candidates, s + 1)
    if e < 0 or e <= s:
        return False

    to_remove = doc.paragraphs[s + 1 : e]
    for p in to_remove:
        remove_paragraph(p)
    return True


def insert_lines_before_heading(doc: Document, end_heading_candidates: list[str], lines: Iterable[str]):
    e = find_any_idx(doc, end_heading_candidates)
    if e < 0:
        return False

    anchor = doc.paragraphs[e]
    for line in lines:
        anchor.insert_paragraph_before(line)
    return True


def replace_section(doc: Document, start_heading: str, end_heading_candidates: list[str], lines: list[str]):
    removed = remove_between(doc, start_heading, end_heading_candidates)
    inserted = insert_lines_before_heading(doc, end_heading_candidates, lines)
    return removed and inserted


def replace_exact(doc: Document, old: str, new: str):
    for p in doc.paragraphs:
        if p_text(p) == old:
            p.text = new
            return True
    return False


def remove_all_embedded_images(doc: Document):
    # Remove paragraphs containing drawing/picture elements.
    to_remove = []
    for p in doc.paragraphs:
        has_drawing = bool(p._element.xpath('.//w:drawing'))
        has_pict = bool(p._element.xpath('.//w:pict'))
        if has_drawing or has_pict:
            to_remove.append(p)

    for p in to_remove:
        remove_paragraph(p)


def insert_diagram_before_heading(doc: Document, end_heading_candidates: list[str], image_path: Path, caption: str):
    e = find_any_idx(doc, end_heading_candidates)
    if e < 0:
        return False

    anchor = doc.paragraphs[e]
    image_para = anchor.insert_paragraph_before("")
    run = image_para.add_run()
    run.add_picture(str(image_path), width=Inches(6.5))
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    caption_para = anchor.insert_paragraph_before(caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return True


def add_screenshot_placeholders(doc: Document):
    replace_section(
        doc,
        "8.1 Input and Output Screenshots",
        ["Facilities Required for Proposed Work"],
        [
            "INPUT AND OUTPUT SCREENSHOTS (SPACE RESERVED)",
            "1) Input Upload Screen - (Insert screenshot here)",
            "",
            "",
            "",
            "2) Prediction Result Screen - (Insert screenshot here)",
            "",
            "",
            "",
            "3) Case Detail Screen - (Insert screenshot here)",
            "",
            "",
            "",
            "4) Patient Report Screen - (Insert screenshot here)",
            "",
            "",
            "",
        ],
    )


def refresh_content(doc: Document):
    # Top matter text corrections only (format stays from template).
    replace_exact(doc, "Project Based Learning - BCA20100", "MINI PROJECT - BCA20090")
    replace_exact(doc, "Date: 1st February 2026", "Date: 24 April 2026")
    replace_exact(doc, "CuraVision AI - AI Powered Medical Image Classification System", "CuraVision AI - AI Powered Medical Image Classification System")

    # Index updates.
    replace_section(
        doc,
        "Index",
        ["Nomenclature/Notations"],
        [
            "1. Nomenclature/Notations",
            "2. Abstract",
            "3. Motivation",
            "4. Literature Review",
            "5. Problem Formulation & Objectives",
            "6. Methodology/ Planning of Work",
            "6.1 Requirement Analysis",
            "6.2 Implementation Phases",
            "6.3 System Architecture",
            "6.4 Database Design",
            "6.5 Database Tables",
            "6.6 Project Flow Diagram",
            "6.7 Testing and Deployment",
            "7. Dataset Specifications",
            "8. Testing and Evaluation",
            "8.1 Input and Output Screenshots",
            "9. Facilities Required for Proposed Work",
            "10. References",
            "",
        ],
    )

    replace_section(
        doc,
        "Nomenclature/Notations",
        ["Abstract"],
        [
            "Artificial Intelligence (AI): Systems that assist decision-making by learning patterns from data.",
            "Convolutional Neural Network (CNN): Deep learning architecture specialized for image feature extraction.",
            "Chest X-ray (CXR): Radiographic image modality used to detect thoracic abnormalities such as pneumonia.",
            "Out-of-Distribution (OOD): Input outside the training domain that can reduce prediction reliability.",
            "Transfer Learning: Reusing pre-trained feature extractors for a related target task.",
            "MobileNetV2: Efficient CNN architecture used for lightweight classification workloads.",
            "Flask: Python web framework used for routing, API orchestration, and template rendering.",
            "TensorFlow/Keras: Libraries used for loading and running the pneumonia classifier model.",
            "Confidence Score: Numeric certainty estimate associated with each prediction result.",
            "Clinical Decision Support System (CDSS): Software that assists clinicians with interpretable computational outputs.",
            "",
        ],
    )

    replace_section(
        doc,
        "Abstract",
        ["Motivation"],
        [
            "CuraVision AI is an educational AI-powered clinical decision support system for chest X-ray triage. The platform combines secure doctor access, two-stage inference, and persistent case storage in one integrated web workflow.",
            "The pipeline first validates image suitability using a chest X-ray gate to reduce out-of-domain misuse. Validated scans are then classified as NORMAL or PNEUMONIA using a TensorFlow model with configurable threshold logic and confidence reporting.",
            "For each case, the system stores patient metadata, prediction output, confidence values, generated symptom text, and timestamped report entries. This supports dashboard analytics, case detail review, and patient identifier based report lookup.",
            "The project is intended for academic use only and all outputs require interpretation by qualified healthcare professionals.",
            "",
        ],
    )

    replace_section(
        doc,
        "Motivation",
        ["Literature Review"],
        [
            "The motivation behind CuraVision AI is to improve reliability and workflow quality in educational medical-AI systems.",
            "Many basic prototypes classify any uploaded image directly, which can produce misleading medical outputs. This project addresses that weakness using a validation-first design.",
            "The work also aims to demonstrate full-stack project competency by combining AI model integration, authentication workflows, persistence, and structured reporting.",
            "Responsible AI usage is a key motivation: prediction confidence and explanatory summaries are shown to support clinician-in-the-loop interpretation.",
            "",
        ],
    )

    replace_section(
        doc,
        "Literature Review",
        ["Problem Formulation & Objectives"],
        [
            "Prior work in chest X-ray classification demonstrates strong performance with transfer learning models such as MobileNetV2.",
            "However, practical reliability requires additional controls for invalid inputs and workflow traceability.",
            "CuraVision AI adopts a two-stage approach that aligns with reviewed best practices: first input suitability validation, then medical classification with confidence outputs.",
            "This design improves robustness for real-world-like usage scenarios in an academic setting.",
            "",
        ],
    )

    replace_section(
        doc,
        "Problem Formulation & Objectives",
        ["Methodology/ Planning of Work"],
        [
            "Problem Statement",
            "Existing educational medical AI systems often lack strong input relevance checks, secure doctor workflow, and persistent report management.",
            "Main Objective",
            "Develop an AI-assisted web platform that validates chest X-ray inputs, predicts NORMAL/PNEUMONIA with confidence, and stores complete case reports.",
            "Specific Objectives",
            "Implement a two-stage inference pipeline with suitability gate and classifier.",
            "Provide confidence-based outputs with interpretable summary text.",
            "Enable doctor authentication and OTP verification for controlled access.",
            "Persist patient report history for dashboard and case review workflows.",
            "Support patient identifier based report retrieval.",
            "",
        ],
    )

    replace_section(
        doc,
        "Methodology/ Planning of Work",
        ["6.1 Requirement Analysis"],
        [
            "The methodology follows a modular implementation plan covering requirement definition, route development, model integration, database persistence, and testing.",
            "A layered design is used to separate authentication, inference, and presentation responsibilities for maintainability and clearer debugging.",
            "The planning phase also includes deployment readiness, startup checks, and report formatting for faculty evaluation.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.1 Requirement Analysis",
        ["6.2 Implementation Phases"],
        [
            "Functional Requirements",
            "Secure Access: Doctor registration, OTP verification, and session-protected routes.",
            "Upload and Validation: Accept chest X-ray inputs and reject invalid formats or non-chest content.",
            "Prediction and Explanation: Return prediction label, confidence score, and AI summary.",
            "Case Persistence: Save complete case records for history and report retrieval.",
            "Software Requirements",
            "Python 3.11, Flask, Flask-SQLAlchemy, Flask-Bcrypt, TensorFlow/Keras, Transformers CLIP, Pillow/OpenCV.",
            "Hardware Requirements",
            "Minimum 8 GB RAM, modern multicore CPU, optional CUDA GPU for faster experiments.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.2 Implementation Phases",
        ["6.3 System Architecture"],
        [
            "System Design: Separate modules for authentication, routes, AI helper logic, and models.",
            "Frontend Development: Dashboard, upload form, case list/detail pages, AI helper, and patient report portal.",
            "Backend Development: Validation-first upload processing and stable API response handling.",
            "Database Integration: SQLAlchemy models for doctors and patients with compatibility checks.",
            "Model Integration: CLIP suitability stage plus TensorFlow pneumonia classifier.",
            "Testing and Hardening: Route-level tests and manual scenario validation for reliability.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.3 System Architecture",
        ["6.4 Testing and Deployment"],
        [
            "Access Layer: Doctor onboarding and authenticated session flow.",
            "Input Layer: X-ray and optional photo upload with patient metadata.",
            "Validation Layer: Suitability check to reject non-chest uploads.",
            "Inference Layer: NORMAL/PNEUMONIA prediction with configurable threshold.",
            "Insight Layer: Rule-based AI summary generation.",
            "Persistence Layer: SQLite storage for reports, media, and timeline metadata.",
            "Presentation Layer: Dashboard, case detail, and patient report retrieval.",
            "",
        ],
    )

    replace_section(
        doc,
        "6.4 Testing and Deployment",
        ["Dataset Specifications", "7. Dataset Specifications"],
        [
            "Development Phases and Deployment",
            "Environment Setup: Virtual environment activation and dependency installation.",
            "Run Strategy: Use pre-trained model files during app startup.",
            "API Integration: Validate upload, prediction, and report retrieval routes end-to-end.",
            "Schema Safety: Startup migration checks for backward compatibility in SQLite.",
            "Deployment Scope: Local academic deployment with responsible-AI disclaimer.",
            "",
        ],
    )

    replace_section(
        doc,
        "Dataset Specifications",
        ["Testing and Evaluation", "8. Testing and Evaluations"],
        [
            "7. Dataset Specifications",
            "Primary Dataset: Kaggle chest X-ray pneumonia data (NORMAL and PNEUMONIA classes).",
            "Runtime Dataset: Application-generated case records stored in SQLite and uploads directory.",
            "Data Attributes: Patient ID, demographics, symptoms, prediction label, confidence, AI insights, and timestamp.",
            "Validation Constraints: Format validation, suitability filtering, and preprocessing consistency.",
            "",
        ],
    )

    replace_section(
        doc,
        "Testing and Evaluation",
        ["Facilities Required for Proposed Work"],
        [
            "8. Testing and Evaluation",
            "Functional Testing: Authentication, upload, prediction rendering, case history, and patient report flow.",
            "Validation Testing: Rejection of invalid uploads and non-chest images with clear error responses.",
            "Database Testing: Record creation, retrieval ordering, and deletion consistency.",
            "Performance Summary (reported): Training 92.64%, Validation 87.66%, Test 87.66%.",
            "Class Metrics (reported): NORMAL F1 81.53%, PNEUMONIA F1 90.73%.",
            "",
            "8.1 Input and Output Screenshots",
        ],
    )

    add_screenshot_placeholders(doc)

    replace_section(
        doc,
        "Facilities Required for Proposed Work",
        ["Bibliography/References"],
        [
            "9. Facilities Required for Proposed Work",
            "9.1 Hardware Requirements",
            "Development machine with minimum 8 GB RAM (16 GB recommended).",
            "Modern multicore CPU and optional CUDA-capable GPU for experimentation.",
            "Adequate storage for datasets, model files, and report backups.",
            "9.2 Software Requirements",
            "Operating System: Windows/macOS/Linux.",
            "IDE and tooling: VS Code, Git, GitHub.",
            "Backend Stack: Python, Flask, Flask-SQLAlchemy, Flask-Bcrypt.",
            "AI Stack: TensorFlow/Keras, Transformers CLIP, Pillow/OpenCV.",
            "Database: SQLite.",
            "",
        ],
    )

    replace_section(
        doc,
        "Bibliography/References",
        ["10. References"],
        [
            "10. References",
            "1. Kaggle Chest X-ray Pneumonia Dataset - https://www.kaggle.com/paultimothymooney/chest-xray-pneumonia",
            "2. TensorFlow Documentation - https://www.tensorflow.org/",
            "3. Keras Documentation - https://keras.io/",
            "4. Flask Documentation - https://flask.palletsprojects.com/",
            "5. Flask-SQLAlchemy Documentation - https://flask-sqlalchemy.palletsprojects.com/",
            "6. Transformers Documentation - https://huggingface.co/docs/transformers/",
            "7. MobileNetV2 Paper - https://arxiv.org/abs/1801.04381",
            "8. SQLite Documentation - https://www.sqlite.org/docs.html",
            "9. Project Repository - https://github.com/vedant420-s/mini-project",
            "",
        ],
    )


def add_our_diagrams(doc: Document):
    # Insert diagrams before the next heading in corresponding sections.
    insert_diagram_before_heading(
        doc,
        ["6.4 Testing and Deployment"],
        DIAGRAMS["system"],
        "Figure: System Architecture Diagram",
    )

    insert_diagram_before_heading(
        doc,
        ["Dataset Specifications", "7. Dataset Specifications"],
        DIAGRAMS["er"],
        "Figure: ER Diagram and Database Models",
    )

    insert_diagram_before_heading(
        doc,
        ["Testing and Evaluation", "8. Testing and Evaluations"],
        DIAGRAMS["flow"],
        "Figure: Project Flow Diagram",
    )

    insert_diagram_before_heading(
        doc,
        ["Facilities Required for Proposed Work"],
        DIAGRAMS["pipeline"],
        "Figure: AI Inference Pipeline",
    )


def main():
    for key, path in DIAGRAMS.items():
        if not path.exists():
            raise FileNotFoundError(f"Missing diagram for {key}: {path}")

    doc = Document(str(SOURCE_DOC))

    remove_all_embedded_images(doc)
    refresh_content(doc)
    add_our_diagrams(doc)

    doc.save(str(OUTPUT_WORKSPACE))
    doc.save(str(OUTPUT_DESKTOP))

    final_doc = Document(str(OUTPUT_WORKSPACE))
    print(f"Updated template saved (workspace): {OUTPUT_WORKSPACE}")
    print(f"Updated template saved (desktop): {OUTPUT_DESKTOP}")
    print(f"Final images in document: {len(final_doc.inline_shapes)}")


if __name__ == "__main__":
    main()
