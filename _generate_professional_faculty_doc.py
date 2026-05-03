from __future__ import annotations

from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"c:\Users\VEDANT\OneDrive\Desktop\mini project")
DESKTOP = Path(r"c:\Users\VEDANT\OneDrive\Desktop")
ASSET_DIR = ROOT / "doc_diagrams" / "professional_faculty"
ASSET_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DOC = DESKTOP / "AI_Medical_Image_Classifier_Faculty_Report.docx"

PROJECT_TITLE = "AI-Powered Clinical Decision Support System for Chest X-ray Triage"
TEAM = "Vedant and Sakshi"
REPO = "https://github.com/vedant420-s/mini-project"


PALETTE = {
    "navy": (17, 45, 78),
    "teal": (24, 128, 122),
    "orange": (223, 124, 47),
    "green": (52, 148, 98),
    "slate": (88, 101, 117),
    "light_blue": (238, 246, 255),
    "light_green": (236, 250, 243),
    "light_orange": (255, 245, 234),
    "light_gray": (246, 247, 250),
    "line": (74, 93, 120),
}


def load_font(size: int, bold: bool = False):
    font_candidates = [
        "segoeuib.ttf" if bold else "segoeui.ttf",
        "calibrib.ttf" if bold else "calibri.ttf",
        "arialbd.ttf" if bold else "arial.ttf",
    ]
    for candidate in font_candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def draw_wrapped_text(draw: ImageDraw.ImageDraw, text: str, box: tuple[int, int, int, int], font, fill=(30, 30, 30), line_gap: int = 6):
    x1, y1, x2, y2 = box
    max_w = x2 - x1

    words = text.split()
    lines: list[str] = []
    line = ""
    for word in words:
        trial = word if not line else f"{line} {word}"
        width = draw.textbbox((0, 0), trial, font=font)[2]
        if width <= max_w:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)

    y = y1
    for ln in lines:
        draw.text((x1, y), ln, font=font, fill=fill)
        h = draw.textbbox((0, 0), ln, font=font)[3]
        y += h + line_gap
        if y > y2:
            break


def draw_card(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], title: str, body: str, border, fill):
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=20, outline=border, width=5, fill=fill)
    title_font = load_font(32, bold=True)
    body_font = load_font(24)
    draw.text((x1 + 22, y1 + 14), title, font=title_font, fill=border)
    draw_wrapped_text(draw, body, (x1 + 22, y1 + 62, x2 - 20, y2 - 18), body_font)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color=(60, 85, 120), width: int = 6):
    sx, sy = start
    ex, ey = end
    draw.line((sx, sy, ex, ey), fill=color, width=width)

    head = 16
    if abs(ex - sx) >= abs(ey - sy):
        if ex >= sx:
            points = [(ex, ey), (ex - head, ey - head // 2), (ex - head, ey + head // 2)]
        else:
            points = [(ex, ey), (ex + head, ey - head // 2), (ex + head, ey + head // 2)]
    else:
        if ey >= sy:
            points = [(ex, ey), (ex - head // 2, ey - head), (ex + head // 2, ey - head)]
        else:
            points = [(ex, ey), (ex - head // 2, ey + head), (ex + head // 2, ey + head)]

    draw.polygon(points, fill=color)


def create_system_architecture_diagram(path: Path):
    img = Image.new("RGB", (2200, 1300), PALETTE["light_gray"])
    draw = ImageDraw.Draw(img)

    title_font = load_font(58, bold=True)
    draw.text((70, 30), "System Architecture Diagram", font=title_font, fill=PALETTE["navy"])

    draw_card(
        draw,
        (90, 180, 650, 500),
        "Presentation Layer",
        "Jinja templates + Tailwind UI pages: Login, Dashboard, Upload, Cases, AI Helper, Patient Portal.",
        PALETTE["navy"],
        PALETTE["light_blue"],
    )
    draw_card(
        draw,
        (820, 180, 1380, 500),
        "Application Layer",
        "Flask routes.py and auth.py handle validation, authentication, workflows, and JSON responses.",
        PALETTE["teal"],
        PALETTE["light_green"],
    )
    draw_card(
        draw,
        (1550, 180, 2110, 500),
        "AI Inference Layer",
        "CLIP chest suitability gate + TensorFlow pneumonia classifier + rule-based AI insights.",
        PALETTE["orange"],
        PALETTE["light_orange"],
    )

    draw_card(
        draw,
        (350, 700, 1030, 1080),
        "Data Layer",
        "SQLite + Flask-SQLAlchemy models: doctors and patients. Stores metadata, predictions, confidence, and image BLOBs.",
        PALETTE["green"],
        (239, 252, 245),
    )
    draw_card(
        draw,
        (1170, 700, 1850, 1080),
        "Integration Services",
        "Email OTP and scan notifications via Resend / SendGrid / SMTP. File storage under uploads/xrays and uploads/photos.",
        PALETTE["slate"],
        (242, 245, 249),
    )

    draw_arrow(draw, (650, 340), (820, 340), PALETTE["line"])
    draw_arrow(draw, (1380, 340), (1550, 340), PALETTE["line"])
    draw_arrow(draw, (1100, 500), (700, 700), PALETTE["line"])
    draw_arrow(draw, (1240, 500), (1500, 700), PALETTE["line"])

    footer_font = load_font(26)
    draw.text((70, 1210), "Layered Architecture for Faculty Demonstration", font=footer_font, fill=PALETTE["slate"])
    img.save(path)


def create_workflow_diagram(path: Path):
    img = Image.new("RGB", (2200, 1400), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    title_font = load_font(56, bold=True)
    draw.text((70, 24), "End-to-End Clinical Workflow", font=title_font, fill=PALETTE["navy"])

    steps = [
        ("1. Doctor Registration", "Sign up, receive OTP, verify email, then login."),
        ("2. Upload Case", "Enter patient details and upload chest X-ray + optional photo."),
        ("3. X-ray Suitability", "CLIP validates that image is a chest X-ray."),
        ("4. Pneumonia Prediction", "TensorFlow model outputs NORMAL/PNEUMONIA + confidence."),
        ("5. AI Insights", "Rule-based layer generates clinical summary text."),
        ("6. Persist & Review", "Store full case and display in dashboard/history pages."),
    ]

    y = 140
    for idx, (title, body) in enumerate(steps):
        top = y + idx * 190
        box = (120, top, 2080, top + 140)
        color_cycle = [PALETTE["navy"], PALETTE["teal"], PALETTE["green"], PALETTE["orange"], PALETTE["slate"], PALETTE["navy"]]
        color = color_cycle[idx]
        draw.rounded_rectangle(box, radius=16, outline=color, width=5, fill=(249, 251, 255))
        draw.text((160, top + 18), title, font=load_font(32, bold=True), fill=color)
        draw.text((160, top + 70), body, font=load_font(26), fill=(40, 40, 40))
        if idx < len(steps) - 1:
            draw_arrow(draw, (1100, top + 140), (1100, top + 190), color=PALETTE["line"], width=6)

    draw.rounded_rectangle((120, 1280, 2080, 1360), radius=16, outline=PALETTE["green"], width=4, fill=(238, 252, 243))
    draw.text(
        (160, 1304),
        "Patient Portal Path: Patient ID -> Reports List -> Report Detail with Confidence and Insights",
        font=load_font(26, bold=True),
        fill=PALETTE["green"],
    )

    img.save(path)


def create_er_diagram(path: Path):
    img = Image.new("RGB", (2200, 1300), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((70, 30), "ER Diagram (Database View)", font=load_font(58, bold=True), fill=PALETTE["navy"])

    left = (120, 180, 960, 1120)
    right = (1240, 180, 2080, 1120)

    draw.rounded_rectangle(left, radius=18, outline=PALETTE["navy"], width=5, fill=PALETTE["light_blue"])
    draw.rounded_rectangle(right, radius=18, outline=PALETTE["teal"], width=5, fill=PALETTE["light_green"])

    draw.text((160, 220), "TABLE: doctors", font=load_font(40, bold=True), fill=PALETTE["navy"])
    draw.text((1280, 220), "TABLE: patients", font=load_font(40, bold=True), fill=PALETTE["teal"])

    doctors_fields = [
        "PK id : Integer",
        "name : String(120)",
        "email : String(255), UNIQUE, INDEX",
        "password_hash : String(255)",
        "is_verified : Boolean",
        "verification_otp : String(10), nullable",
        "otp_expires_at : DateTime, nullable",
    ]

    patients_fields = [
        "PK id : Integer",
        "patient_identifier : String(64), INDEX",
        "name : String(120)",
        "age : Integer",
        "symptoms : Text",
        "generated_symptoms : Text",
        "image_path / photo_path : String(500)",
        "image_blob / photo_blob : LargeBinary",
        "image_mime / photo_mime : String(100)",
        "prediction : String(50)",
        "confidence : Float",
        "ai_insights : Text",
        "created_at : DateTime",
    ]

    y = 300
    for line in doctors_fields:
        draw.text((160, y), line, font=load_font(28), fill=(35, 35, 35))
        y += 58

    y = 300
    for line in patients_fields:
        draw.text((1280, y), line, font=load_font(28), fill=(35, 35, 35))
        y += 58

    draw_arrow(draw, (960, 620), (1240, 620), color=PALETTE["slate"], width=7)
    draw.text((870, 675), "workflow association", font=load_font(25), fill=PALETTE["slate"])

    draw.text(
        (70, 1220),
        "Relational database: SQLite + Flask-SQLAlchemy. Doctor authentication and patient case persistence.",
        font=load_font(25),
        fill=(80, 80, 80),
    )

    img.save(path)


def create_model_diagram(path: Path):
    img = Image.new("RGB", (2200, 1260), PALETTE["light_gray"])
    draw = ImageDraw.Draw(img)
    draw.text((70, 24), "Database Models Class Diagram", font=load_font(58, bold=True), fill=PALETTE["navy"])

    draw.rounded_rectangle((140, 180, 1020, 1100), radius=20, outline=PALETTE["navy"], width=5, fill=(255, 255, 255))
    draw.rounded_rectangle((1180, 180, 2060, 1100), radius=20, outline=PALETTE["teal"], width=5, fill=(255, 255, 255))

    draw.text((180, 220), "class Doctor(db.Model)", font=load_font(36, bold=True), fill=PALETTE["navy"])
    doctor_lines = [
        "__tablename__ = 'doctors'",
        "id: Integer = primary_key",
        "name: String(120)",
        "email: String(255), unique, index",
        "password_hash: String(255)",
        "is_verified: Boolean",
        "verification_otp: String(10) | None",
        "otp_expires_at: DateTime | None",
    ]
    y = 290
    for line in doctor_lines:
        draw.text((180, y), line, font=load_font(28), fill=(45, 45, 45))
        y += 64

    draw.text((1220, 220), "class Patient(db.Model)", font=load_font(36, bold=True), fill=PALETTE["teal"])
    patient_lines = [
        "__tablename__ = 'patients'",
        "id: Integer = primary_key",
        "patient_identifier: String(64)",
        "name: String(120)",
        "age: Integer",
        "symptoms / generated_symptoms: Text",
        "image_path / photo_path: String(500)",
        "image_blob / photo_blob: LargeBinary",
        "prediction: String(50)",
        "confidence: Float",
        "ai_insights: Text",
        "created_at: DateTime",
    ]
    y = 290
    for line in patient_lines:
        draw.text((1220, y), line, font=load_font(28), fill=(45, 45, 45))
        y += 58

    draw_arrow(draw, (1020, 650), (1180, 650), color=PALETTE["line"], width=7)
    draw.text((820, 710), "linked in application workflow", font=load_font(24), fill=PALETTE["slate"])
    img.save(path)


def create_ai_pipeline_diagram(path: Path):
    img = Image.new("RGB", (2200, 1240), (255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((70, 26), "AI Inference Pipeline", font=load_font(58, bold=True), fill=PALETTE["navy"])

    cards = [
        (120, "Input X-ray", "Upload from doctor UI"),
        (470, "Preprocess", "Resize 224x224, normalize pixels"),
        (820, "CLIP Gate", "Reject non-chest images"),
        (1170, "CNN Predict", "NORMAL/PNEUMONIA probability"),
        (1520, "Threshold", "PNEUMONIA_THRESHOLD = 0.40"),
        (1870, "Response", "Prediction + confidence + summary"),
    ]

    for x, title, body in cards:
        draw.rounded_rectangle((x, 340, x + 280, 800), radius=18, outline=PALETTE["navy"], width=4, fill=PALETTE["light_blue"])
        draw.text((x + 20, 390), title, font=load_font(30, bold=True), fill=PALETTE["navy"])
        draw_wrapped_text(draw, body, (x + 20, 470, x + 260, 760), load_font(25), fill=(45, 45, 45), line_gap=8)

    for i in range(len(cards) - 1):
        start_x = cards[i][0] + 280
        end_x = cards[i + 1][0]
        draw_arrow(draw, (start_x, 570), (end_x, 570), color=PALETTE["line"], width=6)

    draw.rounded_rectangle((120, 900, 2150, 1130), radius=18, outline=PALETTE["green"], width=4, fill=(238, 252, 243))
    draw_wrapped_text(
        draw,
        "Post-processing: generate symptom guidance and AI insight summary, persist metadata and media to SQLite, and render dashboard/case detail output.",
        (160, 960, 2110, 1100),
        load_font(29),
        fill=PALETTE["green"],
        line_gap=10,
    )

    img.save(path)


def create_deployment_diagram(path: Path):
    img = Image.new("RGB", (2200, 1300), PALETTE["light_gray"])
    draw = ImageDraw.Draw(img)
    draw.text((70, 24), "Deployment & Runtime Diagram", font=load_font(58, bold=True), fill=PALETTE["navy"])

    draw_card(
        draw,
        (120, 200, 770, 520),
        "Developer Machine",
        "Windows + Python virtual environment + Flask app process (app.py)",
        PALETTE["navy"],
        (255, 255, 255),
    )
    draw_card(
        draw,
        (980, 200, 1620, 520),
        "Runtime Components",
        "TensorFlow model.h5 + CLIP model + routes.py/auth.py + ai_logic.py",
        PALETTE["teal"],
        (255, 255, 255),
    )
    draw_card(
        draw,
        (1760, 200, 2110, 520),
        "External Email",
        "Resend / SendGrid / SMTP",
        PALETTE["orange"],
        (255, 255, 255),
    )

    draw_card(
        draw,
        (280, 710, 980, 1080),
        "Persistent Storage",
        "SQLite instance database + uploads/xrays + uploads/photos",
        PALETTE["green"],
        (255, 255, 255),
    )
    draw_card(
        draw,
        (1180, 710, 1920, 1080),
        "Users",
        "Doctor UI (authenticated) and Patient Portal (ID-based report access)",
        PALETTE["slate"],
        (255, 255, 255),
    )

    draw_arrow(draw, (770, 360), (980, 360), color=PALETTE["line"], width=6)
    draw_arrow(draw, (1620, 360), (1760, 360), color=PALETTE["line"], width=6)
    draw_arrow(draw, (1320, 520), (640, 710), color=PALETTE["line"], width=6)
    draw_arrow(draw, (1320, 520), (1460, 710), color=PALETTE["line"], width=6)

    img.save(path)


def create_screenshot_placeholder(path: Path, title: str):
    img = Image.new("RGB", (1800, 1000), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    border_color = PALETTE["slate"]
    for offset in range(0, 8):
        draw.rounded_rectangle((70 + offset, 70 + offset, 1730 - offset, 930 - offset), radius=24, outline=border_color, width=3)

    draw.text((120, 180), title, font=load_font(48, bold=True), fill=PALETTE["navy"])
    draw.text((120, 290), "Reserved space for working project screenshot", font=load_font(34), fill=(70, 70, 70))
    draw.text((120, 360), "Replace this placeholder with your final input/output image", font=load_font(30), fill=(90, 90, 90))

    draw.line((120, 500, 1680, 500), fill=(130, 130, 130), width=3)
    draw.text((120, 550), "Suggested caption format:", font=load_font(30, bold=True), fill=PALETTE["teal"])
    draw.text((120, 610), "Figure X.X: <Screen Name> - <What is shown>", font=load_font(29), fill=(70, 70, 70))

    img.save(path)


def set_page_border(section):
    sect_pr = section._sectPr
    pg_borders = OxmlElement("w:pgBorders")
    pg_borders.set(qn("w:offsetFrom"), "page")

    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "8")
        elem.set(qn("w:space"), "24")
        elem.set(qn("w:color"), "2F5597")
        pg_borders.append(elem)

    sect_pr.append(pg_borders)


def style_heading(paragraph, color: tuple[int, int, int]):
    if not paragraph.runs:
        paragraph.add_run("")
    for run in paragraph.runs:
        run.font.name = "Cambria"
        run.font.bold = True
        run.font.color.rgb = RGBColor(*color)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.name = "Calibri"
    run.font.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(*PALETTE["slate"])


def generate_diagrams() -> dict[str, Path]:
    files = {
        "system_architecture": ASSET_DIR / "system_architecture.png",
        "workflow": ASSET_DIR / "end_to_end_workflow.png",
        "er": ASSET_DIR / "er_diagram_database.png",
        "models": ASSET_DIR / "database_models_diagram.png",
        "ai_pipeline": ASSET_DIR / "ai_inference_pipeline.png",
        "deployment": ASSET_DIR / "deployment_runtime_diagram.png",
        "placeholder_input": ASSET_DIR / "placeholder_input_screenshot.png",
        "placeholder_output": ASSET_DIR / "placeholder_output_screenshot.png",
    }

    create_system_architecture_diagram(files["system_architecture"])
    create_workflow_diagram(files["workflow"])
    create_er_diagram(files["er"])
    create_model_diagram(files["models"])
    create_ai_pipeline_diagram(files["ai_pipeline"])
    create_deployment_diagram(files["deployment"])
    create_screenshot_placeholder(files["placeholder_input"], "INPUT SCREENSHOT PLACEHOLDER")
    create_screenshot_placeholder(files["placeholder_output"], "OUTPUT SCREENSHOT PLACEHOLDER")
    return files


def build_document(diagrams: dict[str, Path]):
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    set_page_border(section)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    title = doc.add_paragraph("AI MEDICAL IMAGE CLASSIFIER")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.name = "Cambria"
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*PALETTE["navy"])

    subtitle = doc.add_paragraph("Comprehensive Project Documentation for Faculty Evaluation")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.runs[0]
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(*PALETTE["teal"])

    doc.add_paragraph("")

    meta = doc.add_table(rows=5, cols=2)
    meta.style = "Light List Accent 1"
    meta.cell(0, 0).text = "Project Title"
    meta.cell(0, 1).text = PROJECT_TITLE
    meta.cell(1, 0).text = "Team"
    meta.cell(1, 1).text = TEAM
    meta.cell(2, 0).text = "Repository"
    meta.cell(2, 1).text = REPO
    meta.cell(3, 0).text = "Submitted On"
    meta.cell(3, 1).text = datetime.now().strftime("%d %B %Y")
    meta.cell(4, 0).text = "Module"
    meta.cell(4, 1).text = "Mini Project / Final Evaluation"

    for row in meta.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for r in paragraph.runs:
                    r.font.size = Pt(11)

    doc.add_page_break()

    h = doc.add_heading("Index", level=1)
    style_heading(h, PALETTE["navy"])
    index_lines = [
        "1. Abstract",
        "2. Project Objectives",
        "3. Technology Stack",
        "4. System Architecture Diagram",
        "5. End-to-End Workflow Diagram",
        "6. AI Inference Pipeline Diagram",
        "7. ER Diagram and Database Models",
        "8. Functional Modules and Routes",
        "9. Performance and Evaluation Summary",
        "10. Security, Validation, and Reliability",
        "11. Screenshot Placeholders (Input/Output)",
        "12. Conclusion and Future Scope",
        "13. References",
    ]
    for line in index_lines:
        doc.add_paragraph(line)

    h = doc.add_heading("1. Abstract", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_paragraph(
        "This project presents an educational Clinical Decision Support System (CDSS) for chest X-ray triage. "
        "The implemented web application combines doctor authentication, two-stage AI inference, and structured case persistence. "
        "The prediction workflow first validates whether an uploaded image is a chest X-ray and then estimates pneumonia probability "
        "using a TensorFlow model. Outputs are presented with confidence values and AI-generated explanatory insights."
    )

    h = doc.add_heading("2. Project Objectives", level=1)
    style_heading(h, PALETTE["navy"])
    objectives = [
        "Develop a robust AI-assisted workflow for chest X-ray triage (NORMAL vs PNEUMONIA).",
        "Prevent invalid-image misuse through a CLIP-based suitability gate.",
        "Provide a complete doctor workflow: authentication, upload, prediction, and case review.",
        "Store reports and images in a database-backed structure for traceability.",
        "Demonstrate full-stack integration of ML model, backend API, UI templates, and persistence.",
    ]
    for item in objectives:
        doc.add_paragraph(item, style="List Bullet")

    h = doc.add_heading("3. Technology Stack", level=1)
    style_heading(h, PALETTE["navy"])
    stack = doc.add_table(rows=6, cols=2)
    stack.style = "Colorful List Accent 1"
    stack.cell(0, 0).text = "Layer"
    stack.cell(0, 1).text = "Tools / Frameworks"
    stack.cell(1, 0).text = "Backend"
    stack.cell(1, 1).text = "Python, Flask, Flask-SQLAlchemy, Flask-Bcrypt"
    stack.cell(2, 0).text = "AI/ML"
    stack.cell(2, 1).text = "TensorFlow/Keras model.h5, Transformers CLIP"
    stack.cell(3, 0).text = "Frontend"
    stack.cell(3, 1).text = "HTML/Jinja templates, Tailwind-based styling, JavaScript"
    stack.cell(4, 0).text = "Database"
    stack.cell(4, 1).text = "SQLite (instance DB)"
    stack.cell(5, 0).text = "Notifications"
    stack.cell(5, 1).text = "Resend / SendGrid / SMTP email integration"

    h = doc.add_heading("4. System Architecture Diagram", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["system_architecture"]), width=Inches(7.0))
    add_caption(doc, "Figure 4.1: Layer-wise system architecture")

    h = doc.add_heading("5. End-to-End Workflow Diagram", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["workflow"]), width=Inches(7.0))
    add_caption(doc, "Figure 5.1: Doctor and patient report workflow")

    h = doc.add_heading("6. AI Inference Pipeline Diagram", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["ai_pipeline"]), width=Inches(7.0))
    add_caption(doc, "Figure 6.1: Two-stage AI pipeline with confidence-based output")

    h = doc.add_heading("7. ER Diagram and Database Models", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_picture(str(diagrams["er"]), width=Inches(7.0))
    add_caption(doc, "Figure 7.1: ER view of doctors and patients tables")
    doc.add_paragraph("")
    doc.add_picture(str(diagrams["models"]), width=Inches(7.0))
    add_caption(doc, "Figure 7.2: SQLAlchemy class-level database model diagram")

    h = doc.add_heading("8. Functional Modules and Routes", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_paragraph("Major Modules", style="Intense Quote")
    modules = [
        "auth.py: Registration, OTP verification, login/logout, session enforcement.",
        "routes.py: Upload, predict, cases listing/details, patient portal, AI helper endpoints.",
        "models.py: SQLAlchemy models for Doctor and Patient entities.",
        "ai_logic.py: Rule-based insight generation and helper assessment logic.",
        "email_utils.py: OTP and scan alert email dispatch.",
    ]
    for item in modules:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph("Important Routes", style="Intense Quote")
    route_table = doc.add_table(rows=9, cols=2)
    route_table.style = "Medium Grid 1 Accent 2"
    route_table.cell(0, 0).text = "Route"
    route_table.cell(0, 1).text = "Purpose"
    route_table.cell(1, 0).text = "/auth/register"
    route_table.cell(1, 1).text = "Doctor account creation"
    route_table.cell(2, 0).text = "/auth/verify-email"
    route_table.cell(2, 1).text = "OTP verification"
    route_table.cell(3, 0).text = "/dashboard"
    route_table.cell(3, 1).text = "Cases analytics and quick access"
    route_table.cell(4, 0).text = "/upload"
    route_table.cell(4, 1).text = "Case upload form"
    route_table.cell(5, 0).text = "/predict (POST)"
    route_table.cell(5, 1).text = "Main two-stage inference API"
    route_table.cell(6, 0).text = "/cases"
    route_table.cell(6, 1).text = "Case history table"
    route_table.cell(7, 0).text = "/patient/reports/<patient_identifier>"
    route_table.cell(7, 1).text = "Patient self-lookup reports"
    route_table.cell(8, 0).text = "/ai-helper"
    route_table.cell(8, 1).text = "Rule-based support tool"

    h = doc.add_heading("9. Performance and Evaluation Summary", level=1)
    style_heading(h, PALETTE["navy"])
    metrics = doc.add_table(rows=4, cols=2)
    metrics.style = "Colorful Grid Accent 3"
    metrics.cell(0, 0).text = "Metric"
    metrics.cell(0, 1).text = "Reported Value"
    metrics.cell(1, 0).text = "Training Accuracy"
    metrics.cell(1, 1).text = "92.64%"
    metrics.cell(2, 0).text = "Validation Accuracy"
    metrics.cell(2, 1).text = "87.66%"
    metrics.cell(3, 0).text = "Test Accuracy"
    metrics.cell(3, 1).text = "87.66%"

    doc.add_paragraph("Class-wise Highlights", style="Intense Quote")
    doc.add_paragraph("NORMAL: Precision 92.90%, Recall 72.65%, F1 81.53%", style="List Bullet")
    doc.add_paragraph("PNEUMONIA: Precision 85.49%, Recall 96.67%, F1 90.73%", style="List Bullet")

    h = doc.add_heading("10. Security, Validation, and Reliability", level=1)
    style_heading(h, PALETTE["navy"])
    controls = [
        "OTP-based email verification for doctor accounts before access.",
        "Session-protected routes via login_required decorator.",
        "Image extension validation and patient form validation.",
        "Two-stage AI validation to reduce out-of-domain inputs.",
        "Database persistence with rollback handling on failures.",
        "Medical disclaimer enforced: educational support tool, not final diagnosis.",
    ]
    for item in controls:
        doc.add_paragraph(item, style="List Bullet")

    h = doc.add_heading("11. Screenshot Placeholders (Input/Output)", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_paragraph(
        "The following dedicated spaces are intentionally left for final working screenshots. "
        "You can replace each placeholder image with actual project execution outputs before submission."
    )

    doc.add_heading("11.1 Input Screens", level=2)
    doc.add_picture(str(diagrams["placeholder_input"]), width=Inches(6.8))
    add_caption(doc, "Figure 11.1: Login/Upload input screen placeholder")
    doc.add_paragraph("")
    doc.add_picture(str(diagrams["placeholder_input"]), width=Inches(6.8))
    add_caption(doc, "Figure 11.2: Patient detail and symptoms input placeholder")

    doc.add_page_break()

    doc.add_heading("11.2 Output Screens", level=2)
    doc.add_picture(str(diagrams["placeholder_output"]), width=Inches(6.8))
    add_caption(doc, "Figure 11.3: Prediction result output placeholder")
    doc.add_paragraph("")
    doc.add_picture(str(diagrams["placeholder_output"]), width=Inches(6.8))
    add_caption(doc, "Figure 11.4: Dashboard/case-history output placeholder")

    h = doc.add_heading("12. Conclusion and Future Scope", level=1)
    style_heading(h, PALETTE["navy"])
    doc.add_paragraph(
        "This mini project demonstrates a complete practical AI system, combining a deep-learning predictor, "
        "application-level validation, secure authentication, and persistent reporting workflows. "
        "The architecture is suitable for academic evaluation because it shows model integration, software engineering discipline, "
        "and user-focused workflow design in a single deployable implementation."
    )
    doc.add_paragraph("Future enhancements include multi-class disease support, Grad-CAM explainability, cloud deployment, and calibration tuning.")

    h = doc.add_heading("13. References", level=1)
    style_heading(h, PALETTE["navy"])
    refs = [
        "TensorFlow Documentation - https://www.tensorflow.org/",
        "OpenAI CLIP and Transformers Documentation - https://huggingface.co/docs/transformers/",
        "Flask Documentation - https://flask.palletsprojects.com/",
        "Kaggle Chest X-ray Dataset - https://www.kaggle.com/paultimothymooney/chest-xray-pneumonia",
        "Project Repository - https://github.com/vedant420-s/mini-project",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    doc.save(str(OUTPUT_DOC))


def main():
    diagrams = generate_diagrams()
    build_document(diagrams)
    print(f"Generated professional faculty document: {OUTPUT_DOC}")
    print(f"Diagram assets saved in: {ASSET_DIR}")


if __name__ == "__main__":
    main()
