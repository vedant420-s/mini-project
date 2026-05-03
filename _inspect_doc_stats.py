from docx import Document
from pathlib import Path
import zipfile

p = Path(r"C:\Users\VEDANT\OneDrive\Desktop\AI_Medical_Image_Classifier_Faculty_Report.docx")
d = Document(str(p))
text = "\n".join((para.text or "") for para in d.paragraphs)
words = [w for w in text.split() if w.strip()]
with zipfile.ZipFile(p) as zf:
    xml = zf.read("word/document.xml").decode("utf-8", "ignore")

page_breaks = xml.count('w:type="page"')
print("PARAGRAPHS", len(d.paragraphs))
print("TABLES", len(d.tables))
print("WORDS", len(words))
print("EXPLICIT_PAGE_BREAKS", page_breaks)
print("EST_MIN_PAGES", page_breaks + 1)
